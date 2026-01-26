"""
Script para generar los 42 templates .docx para Revisar.ia
Cada template incluye placeholders para personalización multi-tenant
"""
import os
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

PLACEHOLDERS = {
    "empresa": "{{EMPRESA_NOMBRE}}",
    "rfc": "{{EMPRESA_RFC}}",
    "fecha": "{{FECHA_ACTUAL}}",
    "proyecto": "{{PROYECTO_NOMBRE}}",
    "monto": "{{MONTO_CREDITO}}",
    "proveedor": "{{PROVEEDOR_NOMBRE}}",
    "proveedor_rfc": "{{PROVEEDOR_RFC}}",
    "descripcion": "{{DESCRIPCION_SERVICIO}}",
    "folio": "{{FOLIO_PROYECTO}}",
    "responsable": "{{RESPONSABLE_PROYECTO}}",
    "periodo": "{{PERIODO_FISCAL}}",
    "fecha_inicio": "{{FECHA_INICIO}}",
    "fecha_fin": "{{FECHA_FIN}}",
}

def create_header(doc, title, subtitle=None):
    """Add professional header to document"""
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("REVISAR.IA")
    run.bold = True
    run.font.size = Pt(16)
    
    if subtitle:
        doc.add_paragraph(subtitle).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    doc.add_paragraph()

def add_section(doc, title, content):
    """Add a section with title and content"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    
    for line in content:
        doc.add_paragraph(line)
    
    doc.add_paragraph()

def create_footer(doc):
    """Add confidentiality footer"""
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("DOCUMENTO CONFIDENCIAL - " + PLACEHOLDERS["empresa"] + " - " + PLACEHOLDERS["fecha"])
    run.font.size = Pt(8)
    run.italic = True

TEMPLATES = {
    "A1_ESTRATEGIA": [
        {
            "filename": "vision_pilares_estrategicos.docx",
            "title": "VISIÓN Y PILARES ESTRATÉGICOS 2026-2030",
            "sections": [
                ("DATOS DE LA EMPRESA", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"RFC: {PLACEHOLDERS['rfc']}",
                    f"Fecha de Elaboración: {PLACEHOLDERS['fecha']}",
                ]),
                ("VISIÓN ESTRATÉGICA", [
                    "La visión estratégica de la empresa se fundamenta en los siguientes principios:",
                    "• Cumplimiento fiscal integral conforme a la normativa mexicana vigente",
                    "• Transparencia en todas las operaciones comerciales",
                    "• Trazabilidad documental de servicios intangibles",
                    "• Prevención de riesgos fiscales ante auditorías SAT",
                ]),
                ("PILARES ESTRATÉGICOS", [
                    "PILAR 1: RAZÓN DE NEGOCIOS (Art. 5-A CFF)",
                    "• Documentar el beneficio económico esperado de cada operación",
                    "• Justificar la necesidad real del servicio contratado",
                    "",
                    "PILAR 2: MATERIALIDAD (Art. 69-B CFF)",
                    "• Evidenciar la prestación efectiva de servicios",
                    "• Mantener entregables tangibles y verificables",
                    "",
                    "PILAR 3: ESTRICTA INDISPENSABILIDAD (Art. 27 LISR)",
                    "• Demostrar que cada gasto es indispensable para la actividad",
                    "• Vincular gastos con generación de ingresos",
                    "",
                    "PILAR 4: TRAZABILIDAD DOCUMENTAL",
                    "• Mantener expediente completo de cada operación",
                    "• Asegurar cadena de custodia documental",
                ]),
                ("OBJETIVOS ESTRATÉGICOS 2026", [
                    f"Proyecto Actual: {PLACEHOLDERS['proyecto']}",
                    f"Monto Autorizado: {PLACEHOLDERS['monto']} MXN",
                    "• Reducir exposición fiscal en operaciones de servicios intangibles",
                    "• Implementar controles preventivos ante auditorías SAT",
                    "• Digitalizar expedientes de defensa fiscal",
                ]),
            ]
        },
        {
            "filename": "planeacion_okrs.docx",
            "title": "PLANEACIÓN ESTRATÉGICA CON OKRs",
            "sections": [
                ("DATOS GENERALES", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"RFC: {PLACEHOLDERS['rfc']}",
                    f"Período Fiscal: {PLACEHOLDERS['periodo']}",
                ]),
                ("OBJETIVO 1: CUMPLIMIENTO FISCAL", [
                    "Resultado Clave 1.1: 100% de operaciones con documentación de razón de negocios",
                    "Resultado Clave 1.2: Cero observaciones en auditorías SAT",
                    "Resultado Clave 1.3: Expedientes de defensa completos para todas las operaciones >$500,000 MXN",
                ]),
                ("OBJETIVO 2: GESTIÓN DE PROVEEDORES", [
                    "Resultado Clave 2.1: Validar 100% de proveedores contra lista 69-B antes de contratación",
                    "Resultado Clave 2.2: Contratos con cláusulas de materialidad en 100% de servicios intangibles",
                    "Resultado Clave 2.3: Actas de entrega firmadas para todos los servicios",
                ]),
                ("OBJETIVO 3: TRAZABILIDAD DOCUMENTAL", [
                    "Resultado Clave 3.1: Defense File completo en máximo 48 horas post-servicio",
                    "Resultado Clave 3.2: Bitácora digital de todas las deliberaciones internas",
                    "Resultado Clave 3.3: Almacenamiento en nube con respaldo redundante",
                ]),
            ]
        },
        {
            "filename": "panorama_industria.docx",
            "title": "ANÁLISIS DE PANORAMA DE INDUSTRIA",
            "sections": [
                ("CONTEXTO EMPRESARIAL", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"Sector: {{SECTOR_INDUSTRIA}}",
                    f"Fecha de Análisis: {PLACEHOLDERS['fecha']}",
                ]),
                ("CONTEXTO REGULATORIO MEXICANO", [
                    "El entorno fiscal mexicano ha endurecido su postura contra operaciones simuladas:",
                    "• Artículo 69-B CFF: Presunción de operaciones inexistentes",
                    "• Artículo 5-A CFF: Exigencia de razón de negocios",
                    "• Reforma Fiscal 2020-2025: Mayor escrutinio a servicios intangibles",
                ]),
                ("BENCHMARKS DEL SECTOR", [
                    "Prácticas de documentación fiscal en el sector:",
                    "• 78% de empresas AAA mantienen expedientes de defensa preventivos",
                    "• Tiempo promedio de preparación ante auditoría: 15-30 días",
                    "• Costo promedio de defensa fiscal: 3-8% del monto cuestionado",
                ]),
                ("RIESGOS IDENTIFICADOS", [
                    "• Servicios de consultoría sin entregables tangibles",
                    "• Comisiones por intermediación sin soporte documental",
                    "• Honorarios profesionales a personas morales relacionadas",
                ]),
            ]
        },
        {
            "filename": "plan_estrategico.docx",
            "title": "PLAN ESTRATÉGICO DE CUMPLIMIENTO FISCAL",
            "sections": [
                ("INFORMACIÓN DEL PROYECTO", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Responsable: {PLACEHOLDERS['responsable']}",
                    f"Monto: {PLACEHOLDERS['monto']} MXN",
                ]),
                ("FASE 1: PREVENCIÓN (Antes de la contratación)", [
                    "• Validar proveedor en listas 69-B y 69-B Bis del SAT",
                    "• Verificar objeto social compatible con servicio a contratar",
                    "• Solicitar opinión de cumplimiento fiscal vigente",
                    "• Documentar razón de negocios antes de formalizar",
                ]),
                ("FASE 2: FORMALIZACIÓN (Contratación)", [
                    "• Contrato con alcance detallado de servicios",
                    "• Anexos técnicos con entregables específicos",
                    "• Cláusulas de materialidad y evidencia documental",
                    "• Calendario de pagos vinculado a entregables",
                ]),
                ("FASE 3: EJECUCIÓN (Durante el servicio)", [
                    "• Minutas de reuniones de seguimiento",
                    "• Evidencia fotográfica/documental del trabajo",
                    "• Reportes de avance firmados",
                    "• Correspondencia formal por correo electrónico",
                ]),
                ("FASE 4: CIERRE (Post-servicio)", [
                    "• Acta de entrega-recepción de servicios",
                    "• Encuesta de satisfacción documentada",
                    "• Compilación de Defense File completo",
                    "• Archivo en pCloud con estructura estandarizada",
                ]),
            ]
        },
        {
            "filename": "matriz_bee.docx",
            "title": "MATRIZ DE BENEFICIO ECONÓMICO ESPERADO (BEE)",
            "sections": [
                ("DATOS DEL ANÁLISIS", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"Monto de Inversión: {PLACEHOLDERS['monto']} MXN",
                ]),
                ("BENEFICIOS CUANTIFICABLES", [
                    "1. AHORRO DIRECTO:",
                    "   • Reducción de costos operativos: ${{AHORRO_OPERATIVO}} MXN",
                    "   • Eficiencia en procesos: {{PORCENTAJE_EFICIENCIA}}%",
                    "",
                    "2. GENERACIÓN DE INGRESOS:",
                    "   • Incremento proyectado en ventas: ${{INCREMENTO_VENTAS}} MXN",
                    "   • Nuevos clientes/mercados: {{NUEVOS_CLIENTES}}",
                    "",
                    "3. MITIGACIÓN DE RIESGOS:",
                    "   • Reducción de contingencias fiscales: ${{AHORRO_CONTINGENCIAS}} MXN",
                    "   • Cumplimiento normativo: 100%",
                ]),
                ("CÁLCULO DE ROI", [
                    "Inversión Total: " + PLACEHOLDERS['monto'] + " MXN",
                    "Beneficio Esperado Total: ${{BENEFICIO_TOTAL}} MXN",
                    "ROI Proyectado: {{ROI_PORCENTAJE}}%",
                    "Período de Recuperación: {{PERIODO_RECUPERACION}} meses",
                ]),
                ("JUSTIFICACIÓN DE RAZÓN DE NEGOCIOS", [
                    "Conforme al Artículo 5-A del Código Fiscal de la Federación, la presente operación:",
                    "• Genera un beneficio económico cuantificable y demostrable",
                    "• Es indispensable para la operación del negocio",
                    "• No tiene como propósito principal obtener beneficios fiscales",
                    "• Cuenta con sustancia económica real",
                ]),
            ]
        },
    ],
    "A2_PMO": [
        {
            "filename": "poe_fases_f0_f9.docx",
            "title": "PROCEDIMIENTO OPERATIVO ESTÁNDAR - FASES F0-F9",
            "sections": [
                ("INFORMACIÓN GENERAL", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Folio: {PLACEHOLDERS['folio']}",
                ]),
                ("F0: INTAKE - RECEPCIÓN DE SOLICITUD", [
                    "• Recepción de solicitud de servicio vía Wufoo/Email",
                    "• Asignación de folio único al proyecto",
                    "• Validación inicial de información del proveedor",
                    "• Checklist: RFC, Constancia Fiscal, Opinión Cumplimiento",
                ]),
                ("F1: VALIDACIÓN FISCAL", [
                    "• Análisis de razón de negocios (Art. 5-A CFF)",
                    "• Evaluación de materialidad del servicio",
                    "• Verificación contra lista 69-B SAT",
                    "• Dictamen fiscal del proyecto",
                ]),
                ("F2: REVISIÓN FINANCIERA", [
                    "• Análisis presupuestal",
                    "• Cálculo de ROI esperado",
                    "• Validación de límites de autorización",
                    "• Dictamen financiero",
                ]),
                ("F3: REVISIÓN LEGAL", [
                    "• Validación de personalidad jurídica",
                    "• Revisión de estructura contractual",
                    "• Cláusulas de materialidad",
                    "• Dictamen legal",
                ]),
                ("F4-F5: CONSOLIDACIÓN PMO", [
                    "• Integración de dictámenes de todos los agentes",
                    "• Generación de Orden de Compra",
                    "• Preparación de Defense File",
                    "• Envío a Sponsor para aprobación",
                ]),
                ("F6-F9: EJECUCIÓN Y CIERRE", [
                    "• Seguimiento de entregables",
                    "• Recopilación de evidencias",
                    "• Cierre documental",
                    "• Archivo en pCloud",
                ]),
            ]
        },
        {
            "filename": "consolidacion_po.docx",
            "title": "ORDEN DE COMPRA - CONSOLIDACIÓN PMO",
            "sections": [
                ("DATOS DE LA ORDEN", [
                    f"Número de Orden: {PLACEHOLDERS['folio']}-PO",
                    f"Fecha de Emisión: {PLACEHOLDERS['fecha']}",
                    f"Empresa Solicitante: {PLACEHOLDERS['empresa']}",
                    f"RFC: {PLACEHOLDERS['rfc']}",
                ]),
                ("DATOS DEL PROVEEDOR", [
                    f"Razón Social: {PLACEHOLDERS['proveedor']}",
                    f"RFC: {PLACEHOLDERS['proveedor_rfc']}",
                    "Domicilio Fiscal: {{PROVEEDOR_DOMICILIO}}",
                    "Cuenta Bancaria: {{PROVEEDOR_CUENTA}}",
                ]),
                ("DESCRIPCIÓN DEL SERVICIO", [
                    f"Concepto: {PLACEHOLDERS['descripcion']}",
                    f"Monto Total: {PLACEHOLDERS['monto']} MXN (antes de IVA)",
                    "IVA (16%): {{MONTO_IVA}} MXN",
                    "Total con IVA: {{MONTO_TOTAL_IVA}} MXN",
                ]),
                ("CONDICIONES DE PAGO", [
                    "Forma de Pago: Transferencia Electrónica",
                    "Plazo de Pago: 30 días naturales contra factura",
                    "Moneda: Pesos Mexicanos (MXN)",
                ]),
                ("ENTREGABLES REQUERIDOS", [
                    "1. {{ENTREGABLE_1}}",
                    "2. {{ENTREGABLE_2}}",
                    "3. {{ENTREGABLE_3}}",
                    "4. Factura CFDI con complemento de pago",
                ]),
                ("VALIDACIONES REALIZADAS", [
                    "☑ Proveedor validado contra lista 69-B SAT",
                    "☑ Opinión de cumplimiento fiscal vigente",
                    "☑ Razón de negocios documentada",
                    "☑ Dictamen fiscal aprobado",
                    "☑ Dictamen financiero aprobado",
                    "☑ Dictamen legal aprobado",
                ]),
                ("FIRMAS DE AUTORIZACIÓN", [
                    "",
                    "_________________________",
                    "Gerente de PMO",
                    "",
                    "_________________________",
                    "Director de Finanzas",
                    "",
                    "_________________________",
                    "Sponsor del Proyecto",
                ]),
            ]
        },
        {
            "filename": "coordinacion_multiagente.docx",
            "title": "PROTOCOLO DE COORDINACIÓN MULTIAGENTE",
            "sections": [
                ("OBJETIVO", [
                    "Establecer el protocolo de comunicación y coordinación entre los agentes del sistema Revisar.ia para garantizar la correcta validación de proyectos de servicios intangibles.",
                ]),
                ("FLUJO DE COMUNICACIÓN", [
                    "1. PMO → A1_ESTRATEGIA: Solicitud de validación estratégica",
                    "2. A1_ESTRATEGIA → A3_FISCAL: Dictamen estratégico + Solicitud validación fiscal",
                    "3. A3_FISCAL → A5_FINANZAS: Dictamen fiscal + Solicitud validación financiera",
                    "4. A5_FINANZAS → A4_LEGAL: Dictamen financiero + Solicitud validación legal",
                    "5. A4_LEGAL → PMO: Dictamen legal + Consolidación final",
                    "6. PMO → SPONSOR: Expediente completo para aprobación",
                ]),
                ("TIEMPOS DE RESPUESTA (SLA)", [
                    "• Validación Estratégica: 24 horas",
                    "• Validación Fiscal: 48 horas",
                    "• Validación Financiera: 24 horas",
                    "• Validación Legal: 48 horas",
                    "• Consolidación PMO: 24 horas",
                    "• Aprobación Sponsor: 72 horas",
                ]),
                ("ESCALAMIENTO", [
                    "Si un agente excede su SLA:",
                    "• Notificación automática al PMO",
                    "• Escalamiento a Sponsor si excede 24 horas adicionales",
                    "• Registro en bitácora de incidencias",
                ]),
            ]
        },
        {
            "filename": "checklist_tipologia.docx",
            "title": "CHECKLIST DE TIPOLOGÍA DE SERVICIOS",
            "sections": [
                ("DATOS DEL PROYECTO", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"Descripción: {PLACEHOLDERS['descripcion']}",
                ]),
                ("TIPOLOGÍA DEL SERVICIO", [
                    "☐ Servicios Profesionales Independientes",
                    "☐ Consultoría Especializada",
                    "☐ Servicios Técnicos Especializados",
                    "☐ Desarrollo de Software/Tecnología",
                    "☐ Servicios de Gestión y Administración",
                    "☐ Servicios de Capacitación",
                    "☐ Servicios de Marketing/Publicidad",
                    "☐ Otro: {{OTRO_TIPO_SERVICIO}}",
                ]),
                ("NIVEL DE RIESGO FISCAL", [
                    "☐ BAJO - Servicio con entregables tangibles claros",
                    "☐ MEDIO - Servicio intangible con evidencia documental",
                    "☐ ALTO - Servicio intangible sin entregables claros",
                    "☐ CRÍTICO - Operación con partes relacionadas",
                ]),
                ("DOCUMENTACIÓN REQUERIDA POR TIPOLOGÍA", [
                    "Documentos Mínimos:",
                    "☐ Contrato de prestación de servicios",
                    "☐ Anexo técnico con alcance detallado",
                    "☐ Cronograma de entregables",
                    "",
                    "Documentos Adicionales según tipología:",
                    "☐ Reportes de avance (consultoría)",
                    "☐ Código fuente/documentación técnica (software)",
                    "☐ Material de capacitación/listas de asistencia (capacitación)",
                    "☐ Piezas creativas/métricas de campaña (marketing)",
                ]),
            ]
        },
    ],
    "A3_FISCAL": [
        {
            "filename": "razon_negocios_5a.docx",
            "title": "ANÁLISIS DE RAZÓN DE NEGOCIOS - ARTÍCULO 5-A CFF",
            "sections": [
                ("DATOS DEL ANÁLISIS", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"RFC: {PLACEHOLDERS['rfc']}",
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Fecha: {PLACEHOLDERS['fecha']}",
                ]),
                ("FUNDAMENTO LEGAL", [
                    "Artículo 5-A del Código Fiscal de la Federación:",
                    '"Los actos jurídicos que carezcan de una razón de negocios y que generen un beneficio fiscal directo o indirecto, tendrán los efectos fiscales que correspondan a los que se habrían realizado para la obtención del beneficio económico razonablemente esperado por el contribuyente."',
                ]),
                ("ELEMENTOS DE RAZÓN DE NEGOCIOS", [
                    "1. BENEFICIO ECONÓMICO RAZONABLEMENTE ESPERADO:",
                    f"   • Descripción: {PLACEHOLDERS['descripcion']}",
                    "   • Beneficio cuantificado: {{BENEFICIO_CUANTIFICADO}} MXN",
                    "",
                    "2. PROPÓSITO PRINCIPAL NO FISCAL:",
                    "   • La operación tiene un propósito de negocios genuino",
                    "   • No se realiza principalmente para obtener beneficios fiscales",
                    "",
                    "3. SUSTANCIA SOBRE FORMA:",
                    "   • La operación tiene sustancia económica real",
                    "   • Existe una correlación entre la forma jurídica y la realidad económica",
                ]),
                ("DICTAMEN", [
                    "☐ APROBADO - La operación cuenta con razón de negocios válida",
                    "☐ CONDICIONADO - Requiere documentación adicional",
                    "☐ RECHAZADO - No se identifica razón de negocios suficiente",
                    "",
                    "Observaciones:",
                    "{{OBSERVACIONES_RAZON_NEGOCIOS}}",
                ]),
            ]
        },
        {
            "filename": "materialidad_69b.docx",
            "title": "ANÁLISIS DE MATERIALIDAD - ARTÍCULO 69-B CFF",
            "sections": [
                ("DATOS DEL ANÁLISIS", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"RFC Proveedor: {PLACEHOLDERS['proveedor_rfc']}",
                ]),
                ("FUNDAMENTO LEGAL", [
                    "Artículo 69-B del Código Fiscal de la Federación:",
                    "Presunción de inexistencia de operaciones cuando se detecte que un contribuyente ha estado emitiendo comprobantes sin contar con activos, personal, infraestructura o capacidad material.",
                ]),
                ("VERIFICACIÓN DE CAPACIDAD MATERIAL DEL PROVEEDOR", [
                    "☐ Personal: Cuenta con empleados para prestar el servicio",
                    "☐ Infraestructura: Oficinas/instalaciones verificables",
                    "☐ Activos: Equipo y herramientas necesarios",
                    "☐ Experiencia: Historial de servicios similares",
                    "☐ Objeto Social: Compatible con servicio contratado",
                ]),
                ("EVIDENCIA DE MATERIALIDAD DEL SERVICIO", [
                    "☐ Entregables tangibles documentados",
                    "☐ Comunicaciones y correspondencia del proyecto",
                    "☐ Minutas de reuniones de trabajo",
                    "☐ Reportes de avance firmados",
                    "☐ Acta de entrega-recepción",
                    "☐ Evidencia fotográfica/documental",
                ]),
                ("VALIDACIÓN LISTA 69-B", [
                    "Fecha de Consulta: " + PLACEHOLDERS['fecha'],
                    "Resultado:",
                    "☐ LIMPIO - Proveedor NO aparece en lista 69-B",
                    "☐ ALERTA - Proveedor en proceso de aclaración",
                    "☐ RECHAZADO - Proveedor en lista definitiva 69-B",
                ]),
            ]
        },
        {
            "filename": "indispensabilidad_27.docx",
            "title": "ANÁLISIS DE ESTRICTA INDISPENSABILIDAD - ARTÍCULO 27 LISR",
            "sections": [
                ("DATOS DEL ANÁLISIS", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Monto: {PLACEHOLDERS['monto']} MXN",
                ]),
                ("FUNDAMENTO LEGAL", [
                    "Artículo 27 de la Ley del Impuesto Sobre la Renta:",
                    "Las deducciones autorizadas deberán ser estrictamente indispensables para los fines de la actividad del contribuyente.",
                ]),
                ("ANÁLISIS DE INDISPENSABILIDAD", [
                    "1. RELACIÓN CON LA ACTIVIDAD PREPONDERANTE:",
                    f"   • Actividad de la empresa: {{ACTIVIDAD_EMPRESA}}",
                    f"   • Servicio contratado: {PLACEHOLDERS['descripcion']}",
                    "   • Nexo causal: {{NEXO_CAUSAL}}",
                    "",
                    "2. PROPORCIONALIDAD:",
                    f"   • Monto del gasto: {PLACEHOLDERS['monto']} MXN",
                    "   • Porcentaje vs ingresos: {{PORCENTAJE_INGRESOS}}%",
                    "   • Comparativo con mercado: {{COMPARATIVO_MERCADO}}",
                    "",
                    "3. NECESIDAD DEMOSTRABLE:",
                    "   • El servicio es necesario para: {{NECESIDAD_SERVICIO}}",
                    "   • Consecuencia de no contratarlo: {{CONSECUENCIA_NO_CONTRATAR}}",
                ]),
                ("DICTAMEN", [
                    "☐ DEDUCIBLE - Cumple requisitos de estricta indispensabilidad",
                    "☐ PARCIALMENTE DEDUCIBLE - Ver observaciones",
                    "☐ NO DEDUCIBLE - No cumple requisitos",
                ]),
            ]
        },
        {
            "filename": "cff_extractos.docx",
            "title": "EXTRACTOS RELEVANTES DEL CÓDIGO FISCAL DE LA FEDERACIÓN",
            "sections": [
                ("ARTÍCULO 5-A - RAZÓN DE NEGOCIOS", [
                    "Los actos jurídicos que carezcan de una razón de negocios y que generen un beneficio fiscal directo o indirecto, tendrán los efectos fiscales que correspondan a los que se habrían realizado para la obtención del beneficio económico razonablemente esperado por el contribuyente.",
                    "",
                    "Se considera que no existe una razón de negocios cuando el beneficio económico cuantificable presente o futuro sea menor al beneficio fiscal.",
                ]),
                ("ARTÍCULO 69-B - OPERACIONES INEXISTENTES", [
                    "Cuando la autoridad fiscal detecte que un contribuyente ha estado emitiendo comprobantes sin contar con los activos, personal, infraestructura o capacidad material, directa o indirectamente, para prestar los servicios o producir, comercializar o entregar los bienes que amparan tales comprobantes, o bien, que dichos contribuyentes se encuentren no localizados, se presumirá la inexistencia de las operaciones amparadas en tales comprobantes.",
                ]),
                ("ARTÍCULO 29 - REQUISITOS DE COMPROBANTES", [
                    "Cuando las leyes fiscales establezcan la obligación de expedir comprobantes fiscales por los actos o actividades que realicen, por los ingresos que se perciban o por las retenciones de contribuciones que efectúen, los contribuyentes deberán emitirlos mediante documentos digitales a través de la página de Internet del SAT.",
                ]),
                ("APLICACIÓN EN " + PLACEHOLDERS['empresa'], [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    "Artículos aplicables: 5-A, 69-B, 27 LISR",
                    "Cumplimiento verificado: {{CUMPLIMIENTO_CFF}}",
                ]),
            ]
        },
        {
            "filename": "lisr_extractos.docx",
            "title": "EXTRACTOS RELEVANTES DE LA LEY DEL IMPUESTO SOBRE LA RENTA",
            "sections": [
                ("ARTÍCULO 25 - DEDUCCIONES AUTORIZADAS", [
                    "Los contribuyentes podrán efectuar las deducciones siguientes:",
                    "I. Las devoluciones que se reciban o los descuentos o bonificaciones que se hagan en el ejercicio.",
                    "II. El costo de lo vendido.",
                    "III. Los gastos netos de descuentos, bonificaciones o devoluciones.",
                    "IV. Las inversiones.",
                    "...",
                ]),
                ("ARTÍCULO 27 - REQUISITOS DE LAS DEDUCCIONES", [
                    "Las deducciones autorizadas en este Título deberán reunir los siguientes requisitos:",
                    "I. Ser estrictamente indispensables para los fines de la actividad del contribuyente.",
                    "II. Estar amparadas con comprobante fiscal.",
                    "III. Estar registradas en contabilidad.",
                    "IV. Cumplir con las obligaciones en materia de retención y entero de impuestos.",
                    "...",
                ]),
                ("ARTÍCULO 28 - GASTOS NO DEDUCIBLES", [
                    "Para los efectos de este Título, no serán deducibles:",
                    "I. Los pagos por ISR a cargo del propio contribuyente.",
                    "II. Los gastos e inversiones en proporciones respecto de remuneraciones exentas.",
                    "...",
                    "XXXI. Los pagos que no cumplan con el artículo 27, fracción III del CFF.",
                ]),
            ]
        },
        {
            "filename": "casos_sat.docx",
            "title": "CASOS DE AUDITORÍA SAT - ANÁLISIS Y PRECEDENTES",
            "sections": [
                ("CASO 1: SERVICIOS DE CONSULTORÍA SIN ENTREGABLES", [
                    "Situación: Empresa dedujo $5 millones en servicios de consultoría sin reportes tangibles.",
                    "Resolución SAT: Rechazo de deducción por falta de materialidad.",
                    "Lección: Siempre documentar entregables específicos y verificables.",
                ]),
                ("CASO 2: PROVEEDOR EN LISTA 69-B", [
                    "Situación: Empresa contrató proveedor que posteriormente fue incluido en lista 69-B.",
                    "Resolución SAT: Rechazo de deducción + multas.",
                    "Lección: Validar proveedor ANTES de cada operación, no solo al inicio de la relación.",
                ]),
                ("CASO 3: COMISIONES A INTERMEDIARIOS", [
                    "Situación: Pagos por 'gestión comercial' sin evidencia de trabajo realizado.",
                    "Resolución SAT: Rechazo + presunción de operación simulada.",
                    "Lección: Documentar minutas, comunicaciones y resultados de la gestión.",
                ]),
                ("APLICACIÓN AL PROYECTO ACTUAL", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    "Medidas preventivas implementadas:",
                    "• Validación 69-B: {{RESULTADO_69B}}",
                    "• Entregables documentados: {{ENTREGABLES_DOCUMENTADOS}}",
                    "• Evidencia de materialidad: {{EVIDENCIA_MATERIALIDAD}}",
                ]),
            ]
        },
        {
            "filename": "lista_69b.docx",
            "title": "PROTOCOLO DE VERIFICACIÓN LISTA 69-B SAT",
            "sections": [
                ("PROCEDIMIENTO DE VERIFICACIÓN", [
                    "1. Acceder al portal del SAT: https://www.sat.gob.mx",
                    "2. Ir a: Consultas > Comprobantes Fiscales > Lista de contribuyentes con operaciones inexistentes",
                    "3. Ingresar RFC del proveedor a verificar",
                    "4. Documentar resultado con captura de pantalla",
                    "5. Archivar evidencia en Defense File",
                ]),
                ("REGISTRO DE VERIFICACIÓN", [
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"RFC: {PLACEHOLDERS['proveedor_rfc']}",
                    f"Fecha de Consulta: {PLACEHOLDERS['fecha']}",
                    "",
                    "RESULTADO:",
                    "☐ NO LISTADO - Proveedor no aparece en lista 69-B",
                    "☐ EN PROCESO - Proveedor en etapa de aclaración",
                    "☐ DEFINITIVO - Proveedor en lista definitiva (NO CONTRATAR)",
                ]),
                ("FRECUENCIA DE VERIFICACIÓN", [
                    "• Antes de formalizar contrato: OBLIGATORIO",
                    "• Antes de cada pago: RECOMENDADO",
                    "• Mensualmente durante vigencia del contrato: RECOMENDADO",
                ]),
                ("EVIDENCIA DOCUMENTAL", [
                    "Adjuntar a este documento:",
                    "• Captura de pantalla de la consulta",
                    "• Fecha y hora de la consulta",
                    "• Firma del responsable de verificación",
                ]),
            ]
        },
        {
            "filename": "criterios_efos.docx",
            "title": "CRITERIOS DE IDENTIFICACIÓN DE EFOS",
            "sections": [
                ("DEFINICIÓN", [
                    "EFOS: Empresas que Facturan Operaciones Simuladas",
                    "Son contribuyentes que emiten comprobantes fiscales sin tener la capacidad material, técnica o humana para prestar los servicios o entregar los bienes que amparan.",
                ]),
                ("INDICADORES DE ALERTA (RED FLAGS)", [
                    "☐ RFC de reciente creación (menos de 2 años)",
                    "☐ Domicilio fiscal no localizable",
                    "☐ Sin empleados registrados ante IMSS",
                    "☐ Objeto social genérico o incompatible",
                    "☐ Facturación desproporcionada vs capacidad",
                    "☐ Múltiples giros comerciales",
                    "☐ Cambios frecuentes de domicilio fiscal",
                    "☐ Sin activos fijos declarados",
                ]),
                ("VERIFICACIÓN DEL PROVEEDOR", [
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"RFC: {PLACEHOLDERS['proveedor_rfc']}",
                    "",
                    "Verificación de indicadores:",
                    "• Antigüedad del RFC: {{ANTIGUEDAD_RFC}} años",
                    "• Empleados IMSS: {{EMPLEADOS_IMSS}}",
                    "• Domicilio verificado: {{DOMICILIO_VERIFICADO}}",
                    "• Objeto social compatible: {{OBJETO_COMPATIBLE}}",
                ]),
                ("RECOMENDACIÓN", [
                    "☐ PROCEDER - Sin indicadores de alerta",
                    "☐ PRECAUCIÓN - Requiere verificación adicional",
                    "☐ NO PROCEDER - Múltiples indicadores de alerta",
                ]),
            ]
        },
    ],
    "A4_LEGAL": [
        {
            "filename": "plantilla_contrato.docx",
            "title": "CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES",
            "sections": [
                ("ENCABEZADO", [
                    f"Contrato que celebran:",
                    f"Por una parte: {PLACEHOLDERS['empresa']}, con RFC {PLACEHOLDERS['rfc']} (en adelante 'EL CLIENTE')",
                    f"Por otra parte: {PLACEHOLDERS['proveedor']}, con RFC {PLACEHOLDERS['proveedor_rfc']} (en adelante 'EL PRESTADOR')",
                ]),
                ("DECLARACIONES", [
                    "I. Declara EL CLIENTE:",
                    "   a) Ser una sociedad legalmente constituida conforme a las leyes mexicanas.",
                    "   b) Tener la capacidad legal para celebrar el presente contrato.",
                    f"   c) Su domicilio fiscal es: {{CLIENTE_DOMICILIO}}",
                    "",
                    "II. Declara EL PRESTADOR:",
                    "   a) Ser una persona moral/física con capacidad legal para prestar servicios.",
                    "   b) Contar con los recursos técnicos, humanos y materiales para el servicio.",
                    f"   c) Su domicilio fiscal es: {{PRESTADOR_DOMICILIO}}",
                ]),
                ("CLÁUSULA PRIMERA - OBJETO DEL CONTRATO", [
                    f"EL PRESTADOR se obliga a proporcionar a EL CLIENTE los siguientes servicios:",
                    f"{PLACEHOLDERS['descripcion']}",
                    "",
                    "Los servicios comprenden específicamente:",
                    "a) {{SERVICIO_DETALLE_1}}",
                    "b) {{SERVICIO_DETALLE_2}}",
                    "c) {{SERVICIO_DETALLE_3}}",
                ]),
                ("CLÁUSULA SEGUNDA - CONTRAPRESTACIÓN", [
                    f"Por los servicios objeto de este contrato, EL CLIENTE pagará a EL PRESTADOR:",
                    f"Monto: {PLACEHOLDERS['monto']} MXN (antes de IVA)",
                    "Forma de pago: Transferencia electrónica",
                    "Plazo de pago: 30 días naturales contra factura y entregables aprobados",
                ]),
                ("CLÁUSULA TERCERA - VIGENCIA", [
                    f"Fecha de inicio: {PLACEHOLDERS['fecha_inicio']}",
                    f"Fecha de término: {PLACEHOLDERS['fecha_fin']}",
                ]),
                ("CLÁUSULA CUARTA - ENTREGABLES Y MATERIALIDAD", [
                    "EL PRESTADOR se compromete a entregar los siguientes entregables verificables:",
                    "1. {{ENTREGABLE_1}} - Fecha: {{FECHA_ENTREGABLE_1}}",
                    "2. {{ENTREGABLE_2}} - Fecha: {{FECHA_ENTREGABLE_2}}",
                    "3. {{ENTREGABLE_3}} - Fecha: {{FECHA_ENTREGABLE_3}}",
                    "",
                    "Cada entregable deberá ser documentado y aprobado mediante Acta de Entrega-Recepción.",
                ]),
                ("CLÁUSULA QUINTA - CUMPLIMIENTO FISCAL", [
                    "EL PRESTADOR declara:",
                    "a) No encontrarse en los supuestos del artículo 69-B del CFF",
                    "b) Contar con opinión de cumplimiento de obligaciones fiscales vigente y positiva",
                    "c) Estar al corriente en sus obligaciones fiscales y de seguridad social",
                ]),
                ("FIRMAS", [
                    "",
                    f"_________________________          _________________________",
                    f"    {PLACEHOLDERS['empresa']}               {PLACEHOLDERS['proveedor']}",
                    "         (EL CLIENTE)                    (EL PRESTADOR)",
                    "",
                    f"Fecha: {PLACEHOLDERS['fecha']}",
                ]),
            ]
        },
        {
            "filename": "nom_151_digitalización.docx",
            "title": "PROTOCOLO DE DIGITALIZACIÓN - NOM-151-SCFI-2016",
            "sections": [
                ("OBJETIVO", [
                    "Establecer los requisitos para la digitalización y conservación de documentos que deberán cumplir los sistemas de administración de mensajes de datos, conforme a la NOM-151-SCFI-2016.",
                ]),
                ("REQUISITOS DE DIGITALIZACIÓN", [
                    "1. INTEGRIDAD: Los documentos digitalizados deben ser íntegros y completos",
                    "2. AUTENTICIDAD: Deben poder verificarse el origen y autor del documento",
                    "3. NO REPUDIO: El emisor no puede negar haber generado el documento",
                    "4. CONFIDENCIALIDAD: Acceso controlado y protección de información",
                ]),
                ("PROCESO DE DIGITALIZACIÓN", [
                    "Paso 1: Escaneo del documento original en formato PDF/A",
                    "Paso 2: Generación de huella digital (hash SHA-256)",
                    "Paso 3: Sellado de tiempo con Prestador de Servicios de Certificación",
                    "Paso 4: Almacenamiento en sistema seguro con respaldo",
                    "Paso 5: Registro en bitácora de digitalización",
                ]),
                ("REGISTRO DE DIGITALIZACIÓN", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    "Documentos digitalizados:",
                    "☐ Contrato firmado",
                    "☐ Anexos técnicos",
                    "☐ Actas de entrega",
                    "☐ Facturas CFDI",
                    "☐ Comprobantes de pago",
                ]),
            ]
        },
        {
            "filename": "clausulas_materialidad.docx",
            "title": "CLÁUSULAS CONTRACTUALES DE MATERIALIDAD",
            "sections": [
                ("CLÁUSULA DE MATERIALIDAD DEL SERVICIO", [
                    "EL PRESTADOR se obliga a:",
                    "a) Documentar cada etapa del servicio prestado",
                    "b) Proporcionar evidencia tangible de los trabajos realizados",
                    "c) Entregar reportes de avance mensuales/semanales",
                    "d) Permitir la verificación in situ de los servicios",
                    "e) Mantener registro de comunicaciones y reuniones",
                ]),
                ("CLÁUSULA DE ENTREGABLES VERIFICABLES", [
                    "Los entregables mínimos obligatorios son:",
                    "a) Reportes técnicos con análisis y recomendaciones",
                    "b) Minutas de reuniones firmadas por ambas partes",
                    "c) Correspondencia electrónica del proyecto",
                    "d) Acta de entrega-recepción por cada entregable",
                    "e) Evidencia fotográfica o documental según aplique",
                ]),
                ("CLÁUSULA DE CUMPLIMIENTO FISCAL", [
                    "EL PRESTADOR declara y garantiza:",
                    "a) No estar incluido en las listas publicadas conforme al artículo 69-B del CFF",
                    "b) Contar con opinión de cumplimiento de obligaciones fiscales positiva",
                    "c) Notificar inmediatamente si cambia su situación fiscal",
                    "d) Indemnizar a EL CLIENTE por cualquier contingencia fiscal derivada de incumplimiento",
                ]),
                ("CLÁUSULA DE AUDITORÍA", [
                    "EL CLIENTE tendrá derecho a:",
                    "a) Auditar los registros relacionados con los servicios prestados",
                    "b) Solicitar documentación soporte adicional",
                    "c) Verificar la capacidad operativa del PRESTADOR",
                    "d) Requerir aclaraciones sobre cualquier aspecto del servicio",
                ]),
            ]
        },
        {
            "filename": "sow_alcance_servicios.docx",
            "title": "STATEMENT OF WORK (SOW) - ALCANCE DE SERVICIOS",
            "sections": [
                ("INFORMACIÓN GENERAL", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Cliente: {PLACEHOLDERS['empresa']}",
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"Folio: {PLACEHOLDERS['folio']}",
                ]),
                ("1. ALCANCE DEL PROYECTO", [
                    "1.1 Objetivo General:",
                    f"    {PLACEHOLDERS['descripcion']}",
                    "",
                    "1.2 Objetivos Específicos:",
                    "    a) {{OBJETIVO_1}}",
                    "    b) {{OBJETIVO_2}}",
                    "    c) {{OBJETIVO_3}}",
                    "",
                    "1.3 Fuera de Alcance:",
                    "    a) {{FUERA_ALCANCE_1}}",
                    "    b) {{FUERA_ALCANCE_2}}",
                ]),
                ("2. ENTREGABLES", [
                    "| # | Entregable | Descripción | Fecha | Criterio Aceptación |",
                    "|---|------------|-------------|-------|---------------------|",
                    "| 1 | {{E1_NOMBRE}} | {{E1_DESC}} | {{E1_FECHA}} | {{E1_CRITERIO}} |",
                    "| 2 | {{E2_NOMBRE}} | {{E2_DESC}} | {{E2_FECHA}} | {{E2_CRITERIO}} |",
                    "| 3 | {{E3_NOMBRE}} | {{E3_DESC}} | {{E3_FECHA}} | {{E3_CRITERIO}} |",
                ]),
                ("3. CRONOGRAMA", [
                    f"Fecha de Inicio: {PLACEHOLDERS['fecha_inicio']}",
                    f"Fecha de Fin: {PLACEHOLDERS['fecha_fin']}",
                    "",
                    "Hitos principales:",
                    "• Kick-off: {{FECHA_KICKOFF}}",
                    "• Revisión intermedia: {{FECHA_REVISION}}",
                    "• Entrega final: {{FECHA_ENTREGA_FINAL}}",
                ]),
                ("4. RECURSOS Y RESPONSABILIDADES", [
                    "Por parte del PROVEEDOR:",
                    "• {{RECURSO_PROVEEDOR_1}}",
                    "• {{RECURSO_PROVEEDOR_2}}",
                    "",
                    "Por parte del CLIENTE:",
                    "• {{RECURSO_CLIENTE_1}}",
                    "• {{RECURSO_CLIENTE_2}}",
                ]),
            ]
        },
        {
            "filename": "pagos_penalizaciones.docx",
            "title": "ANEXO DE TÉRMINOS DE PAGO Y PENALIZACIONES",
            "sections": [
                ("DATOS DEL CONTRATO", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Monto Total: {PLACEHOLDERS['monto']} MXN",
                ]),
                ("ESQUEMA DE PAGOS", [
                    "Los pagos se realizarán conforme al siguiente esquema:",
                    "",
                    "| Hito | Concepto | Porcentaje | Monto | Fecha Límite |",
                    "|------|----------|------------|-------|--------------|",
                    "| 1 | Anticipo contra firma | 30% | {{MONTO_ANTICIPO}} | {{FECHA_ANTICIPO}} |",
                    "| 2 | Avance 50% | 30% | {{MONTO_AVANCE}} | {{FECHA_AVANCE}} |",
                    "| 3 | Entrega final | 40% | {{MONTO_FINAL}} | {{FECHA_FINAL}} |",
                ]),
                ("REQUISITOS PARA PAGO", [
                    "Para liberar cada pago se requiere:",
                    "☐ Entregable aprobado por el Cliente",
                    "☐ Acta de entrega-recepción firmada",
                    "☐ Factura CFDI válida y timbrada",
                    "☐ Opinión de cumplimiento fiscal vigente",
                ]),
                ("PENALIZACIONES POR INCUMPLIMIENTO", [
                    "En caso de retraso en entregables:",
                    "• 1-5 días: Penalización del 1% del valor del entregable por día",
                    "• 6-10 días: Penalización del 2% del valor del entregable por día",
                    "• Más de 10 días: Derecho a rescindir contrato y retener garantía",
                    "",
                    "Límite máximo de penalizaciones: 10% del valor total del contrato",
                ]),
            ]
        },
        {
            "filename": "proteccion_datos.docx",
            "title": "AVISO DE PRIVACIDAD Y PROTECCIÓN DE DATOS",
            "sections": [
                ("RESPONSABLE DEL TRATAMIENTO", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"RFC: {PLACEHOLDERS['rfc']}",
                    "Domicilio: {{DOMICILIO_EMPRESA}}",
                    "Correo de contacto: {{EMAIL_PRIVACIDAD}}",
                ]),
                ("FINALIDAD DEL TRATAMIENTO", [
                    "Los datos personales recabados serán utilizados para:",
                    "a) Formalización y cumplimiento del contrato de servicios",
                    "b) Verificación de identidad y capacidad legal",
                    "c) Cumplimiento de obligaciones fiscales",
                    "d) Comunicación relacionada con el proyecto",
                ]),
                ("DATOS PERSONALES RECABADOS", [
                    "☐ Nombre completo",
                    "☐ RFC",
                    "☐ Domicilio fiscal",
                    "☐ Correo electrónico",
                    "☐ Teléfono de contacto",
                    "☐ Datos bancarios para pagos",
                ]),
                ("DERECHOS ARCO", [
                    "El titular tiene derecho a:",
                    "• ACCESO a sus datos personales",
                    "• RECTIFICACIÓN de datos inexactos",
                    "• CANCELACIÓN de datos",
                    "• OPOSICIÓN al tratamiento",
                    "",
                    "Para ejercer estos derechos, contactar a: {{EMAIL_PRIVACIDAD}}",
                ]),
                ("CONSENTIMIENTO", [
                    "",
                    "Acepto el presente Aviso de Privacidad:",
                    "",
                    "_________________________",
                    f"{PLACEHOLDERS['proveedor']}",
                    "",
                    f"Fecha: {PLACEHOLDERS['fecha']}",
                ]),
            ]
        },
    ],
    "A5_FINANZAS": [
        {
            "filename": "politicas_presupuestales.docx",
            "title": "POLÍTICAS PRESUPUESTALES 2026",
            "sections": [
                ("DATOS GENERALES", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"Período Fiscal: {PLACEHOLDERS['periodo']}",
                    f"Fecha de Vigencia: {PLACEHOLDERS['fecha']}",
                ]),
                ("LÍMITES DE AUTORIZACIÓN", [
                    "Nivel 1 - Gerente de Área:",
                    "• Hasta $100,000 MXN por operación",
                    "• Requiere: Cotización y justificación",
                    "",
                    "Nivel 2 - Director de Área:",
                    "• De $100,001 a $500,000 MXN",
                    "• Requiere: 3 cotizaciones + validación fiscal",
                    "",
                    "Nivel 3 - Director General:",
                    "• De $500,001 a $2,000,000 MXN",
                    "• Requiere: Expediente completo + dictámenes",
                    "",
                    "Nivel 4 - Consejo de Administración:",
                    "• Más de $2,000,000 MXN",
                    "• Requiere: Presentación ejecutiva + aprobación en sesión",
                ]),
                ("PARTIDAS PRESUPUESTALES", [
                    "1000 - Servicios Personales",
                    "2000 - Materiales y Suministros",
                    "3000 - Servicios Generales (incluye consultoría)",
                    "4000 - Transferencias",
                    "5000 - Bienes Muebles e Inmuebles",
                    "6000 - Inversiones",
                ]),
                ("REQUISITOS PARA SERVICIOS INTANGIBLES (3000)", [
                    "Todo gasto en servicios intangibles requiere:",
                    "☐ Justificación de razón de negocios",
                    "☐ Análisis de beneficio económico esperado",
                    "☐ Validación de proveedor (69-B)",
                    "☐ Contrato con cláusulas de materialidad",
                ]),
            ]
        },
        {
            "filename": "benchmarks_roi.docx",
            "title": "BENCHMARKS DE ROI POR TIPO DE SERVICIO",
            "sections": [
                ("SERVICIOS DE CONSULTORÍA ESTRATÉGICA", [
                    "ROI esperado: 150-300%",
                    "Período de recuperación: 6-12 meses",
                    "Indicadores clave: Incremento en ingresos, reducción de costos",
                ]),
                ("SERVICIOS DE TECNOLOGÍA/SOFTWARE", [
                    "ROI esperado: 200-400%",
                    "Período de recuperación: 12-24 meses",
                    "Indicadores clave: Automatización, eficiencia operativa",
                ]),
                ("SERVICIOS DE CAPACITACIÓN", [
                    "ROI esperado: 100-200%",
                    "Período de recuperación: 6-18 meses",
                    "Indicadores clave: Productividad, reducción de errores",
                ]),
                ("SERVICIOS FISCALES/LEGALES", [
                    "ROI esperado: 300-500%",
                    "Período de recuperación: 3-6 meses",
                    "Indicadores clave: Ahorro en contingencias, cumplimiento",
                ]),
                ("APLICACIÓN AL PROYECTO", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Inversión: {PLACEHOLDERS['monto']} MXN",
                    "Categoría: {{CATEGORIA_SERVICIO}}",
                    "ROI Proyectado: {{ROI_PROYECTADO}}%",
                    "Período Recuperación: {{PERIODO_RECUPERACION}} meses",
                ]),
            ]
        },
        {
            "filename": "three_way_match.docx",
            "title": "VALIDACIÓN 3-WAY MATCH",
            "sections": [
                ("CONCEPTO", [
                    "El 3-Way Match es un control interno que valida la consistencia entre:",
                    "1. ORDEN DE COMPRA (PO)",
                    "2. RECEPCIÓN/ENTREGA",
                    "3. FACTURA",
                ]),
                ("REGISTRO DE VALIDACIÓN", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"Monto: {PLACEHOLDERS['monto']} MXN",
                ]),
                ("VALIDACIÓN DOCUMENTO 1: ORDEN DE COMPRA", [
                    f"Número de PO: {PLACEHOLDERS['folio']}-PO",
                    "Monto autorizado: {{MONTO_PO}} MXN",
                    "Fecha de emisión: {{FECHA_PO}}",
                    "Firmada por: {{FIRMA_PO}}",
                    "☐ VERIFICADO",
                ]),
                ("VALIDACIÓN DOCUMENTO 2: RECEPCIÓN", [
                    "Acta de Entrega: {{NUMERO_ACTA}}",
                    "Fecha de recepción: {{FECHA_RECEPCION}}",
                    "Recibido por: {{RECIBIO}}",
                    "Entregables recibidos:",
                    "☐ {{ENTREGABLE_1}}",
                    "☐ {{ENTREGABLE_2}}",
                    "☐ VERIFICADO",
                ]),
                ("VALIDACIÓN DOCUMENTO 3: FACTURA", [
                    "UUID Factura: {{UUID_FACTURA}}",
                    "Monto facturado: {{MONTO_FACTURA}} MXN",
                    "Fecha de emisión: {{FECHA_FACTURA}}",
                    "☐ Monto coincide con PO",
                    "☐ Concepto coincide con servicio recibido",
                    "☐ RFC correcto",
                    "☐ VERIFICADO",
                ]),
                ("RESULTADO", [
                    "☐ APROBADO - Los 3 documentos coinciden",
                    "☐ RECHAZADO - Discrepancias detectadas",
                    "",
                    "Discrepancias (si aplica):",
                    "{{DISCREPANCIAS}}",
                ]),
            ]
        },
        {
            "filename": "analisis_financiero.docx",
            "title": "ANÁLISIS FINANCIERO DEL PROYECTO",
            "sections": [
                ("DATOS DEL PROYECTO", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"Monto de Inversión: {PLACEHOLDERS['monto']} MXN",
                ]),
                ("ANÁLISIS DE VIABILIDAD", [
                    "1. DISPONIBILIDAD PRESUPUESTAL:",
                    "   • Partida: 3000 - Servicios Generales",
                    "   • Presupuesto asignado: {{PRESUPUESTO_PARTIDA}} MXN",
                    "   • Comprometido: {{COMPROMETIDO}} MXN",
                    "   • Disponible: {{DISPONIBLE}} MXN",
                    f"   • Requerido: {PLACEHOLDERS['monto']} MXN",
                    "",
                    "2. ANÁLISIS COSTO-BENEFICIO:",
                    "   • Beneficio esperado: {{BENEFICIO_ESPERADO}} MXN",
                    "   • Relación B/C: {{RELACION_BC}}",
                ]),
                ("CÁLCULO DE ROI", [
                    "Inversión Total: " + PLACEHOLDERS['monto'] + " MXN",
                    "Beneficios Proyectados:",
                    "  • Año 1: {{BENEFICIO_ANO_1}} MXN",
                    "  • Año 2: {{BENEFICIO_ANO_2}} MXN",
                    "  • Año 3: {{BENEFICIO_ANO_3}} MXN",
                    "",
                    "ROI = (Beneficio - Inversión) / Inversión × 100",
                    "ROI Proyectado: {{ROI_CALCULADO}}%",
                ]),
                ("DICTAMEN FINANCIERO", [
                    "☐ APROBADO - Viable financieramente",
                    "☐ CONDICIONADO - Ver observaciones",
                    "☐ RECHAZADO - No viable",
                    "",
                    "Observaciones: {{OBSERVACIONES_FINANCIERAS}}",
                ]),
            ]
        },
        {
            "filename": "limites_autorizacion.docx",
            "title": "MATRIZ DE LÍMITES DE AUTORIZACIÓN",
            "sections": [
                ("ESTRUCTURA DE AUTORIZACIONES", [
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"Vigencia: {PLACEHOLDERS['fecha']}",
                ]),
                ("NIVEL 1: OPERATIVO", [
                    "Monto máximo: $50,000 MXN",
                    "Autoriza: Coordinador de Área",
                    "Documentación: Solicitud + Cotización",
                    "Tiempo de respuesta: 24 horas",
                ]),
                ("NIVEL 2: TÁCTICO", [
                    "Monto: $50,001 - $250,000 MXN",
                    "Autoriza: Gerente de Área",
                    "Documentación: Solicitud + 2 Cotizaciones + Justificación",
                    "Tiempo de respuesta: 48 horas",
                ]),
                ("NIVEL 3: ESTRATÉGICO", [
                    "Monto: $250,001 - $1,000,000 MXN",
                    "Autoriza: Director de Área",
                    "Documentación: Expediente completo + Dictámenes",
                    "Tiempo de respuesta: 72 horas",
                    "Validaciones obligatorias: Fiscal, Legal, Financiero",
                ]),
                ("NIVEL 4: EJECUTIVO", [
                    "Monto: Mayor a $1,000,000 MXN",
                    "Autoriza: Director General / Consejo",
                    "Documentación: Expediente completo + Presentación ejecutiva",
                    "Tiempo de respuesta: 5 días hábiles",
                    "Validaciones obligatorias: Todas + Defense File preventivo",
                ]),
                ("REGISTRO DE AUTORIZACIÓN", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Monto: {PLACEHOLDERS['monto']} MXN",
                    "Nivel requerido: {{NIVEL_AUTORIZACION}}",
                    "Autorizado por: {{AUTORIZADOR}}",
                    "Fecha: {{FECHA_AUTORIZACION}}",
                ]),
            ]
        },
    ],
    "A6_PROVEEDOR": [
        {
            "filename": "guia_entregables.docx",
            "title": "GUÍA DE ENTREGABLES PARA PROVEEDORES",
            "sections": [
                ("OBJETIVO", [
                    "Esta guía establece los requisitos mínimos de documentación que todo proveedor de servicios debe cumplir para garantizar la materialidad y trazabilidad de los servicios prestados.",
                ]),
                ("DOCUMENTACIÓN OBLIGATORIA", [
                    "1. ANTES DEL INICIO DEL SERVICIO:",
                    "   ☐ Constancia de Situación Fiscal vigente",
                    "   ☐ Opinión de Cumplimiento SAT positiva",
                    "   ☐ Acta constitutiva (personas morales)",
                    "   ☐ Poder del representante legal",
                    "",
                    "2. DURANTE LA EJECUCIÓN:",
                    "   ☐ Reportes de avance semanales/mensuales",
                    "   ☐ Minutas de reuniones firmadas",
                    "   ☐ Evidencia documental del trabajo (capturas, reportes, etc.)",
                    "",
                    "3. AL CIERRE DEL SERVICIO:",
                    "   ☐ Entregables finales según contrato",
                    "   ☐ Acta de entrega-recepción firmada",
                    "   ☐ Factura CFDI con conceptos detallados",
                ]),
                ("FORMATO DE REPORTES DE AVANCE", [
                    "Cada reporte debe incluir:",
                    "• Período que cubre",
                    "• Actividades realizadas con detalle",
                    "• Horas invertidas",
                    "• Avance vs plan original (%)",
                    "• Próximos pasos",
                    "• Obstáculos o riesgos identificados",
                ]),
                ("INFORMACIÓN DEL PROVEEDOR", [
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"RFC: {PLACEHOLDERS['proveedor_rfc']}",
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Cliente: {PLACEHOLDERS['empresa']}",
                ]),
            ]
        },
        {
            "filename": "acta_aceptacion.docx",
            "title": "ACTA DE ENTREGA-RECEPCIÓN DE SERVICIOS",
            "sections": [
                ("DATOS GENERALES", [
                    f"Fecha: {PLACEHOLDERS['fecha']}",
                    f"Lugar: {{LUGAR_ENTREGA}}",
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Folio: {PLACEHOLDERS['folio']}",
                ]),
                ("PARTES INVOLUCRADAS", [
                    f"ENTREGA: {PLACEHOLDERS['proveedor']} (EL PRESTADOR)",
                    "Representante: {{REPRESENTANTE_PROVEEDOR}}",
                    "Cargo: {{CARGO_PROVEEDOR}}",
                    "",
                    f"RECIBE: {PLACEHOLDERS['empresa']} (EL CLIENTE)",
                    "Representante: {{REPRESENTANTE_CLIENTE}}",
                    "Cargo: {{CARGO_CLIENTE}}",
                ]),
                ("ENTREGABLES RECIBIDOS", [
                    "| # | Entregable | Descripción | Conforme |",
                    "|---|------------|-------------|----------|",
                    "| 1 | {{E1_NOMBRE}} | {{E1_DESC}} | ☐ Sí ☐ No |",
                    "| 2 | {{E2_NOMBRE}} | {{E2_DESC}} | ☐ Sí ☐ No |",
                    "| 3 | {{E3_NOMBRE}} | {{E3_DESC}} | ☐ Sí ☐ No |",
                ]),
                ("VERIFICACIÓN DE CUMPLIMIENTO", [
                    "☐ Los entregables cumplen con las especificaciones del contrato",
                    "☐ La calidad es satisfactoria",
                    "☐ Se entregó en tiempo y forma",
                    "☐ No existen pendientes por parte del PRESTADOR",
                ]),
                ("OBSERVACIONES", [
                    "{{OBSERVACIONES_ACTA}}",
                ]),
                ("FIRMAS", [
                    "",
                    "_________________________          _________________________",
                    "         ENTREGA                            RECIBE",
                    f"   {PLACEHOLDERS['proveedor']}               {PLACEHOLDERS['empresa']}",
                ]),
            ]
        },
        {
            "filename": "checklist_evidencia.docx",
            "title": "CHECKLIST DE EVIDENCIA DOCUMENTAL",
            "sections": [
                ("PROYECTO", [
                    f"Nombre: {PLACEHOLDERS['proyecto']}",
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"Monto: {PLACEHOLDERS['monto']} MXN",
                    f"Período: {PLACEHOLDERS['fecha_inicio']} al {PLACEHOLDERS['fecha_fin']}",
                ]),
                ("DOCUMENTACIÓN FISCAL (OBLIGATORIA)", [
                    "☐ Constancia de Situación Fiscal del proveedor",
                    "☐ Opinión de Cumplimiento SAT (positiva y vigente)",
                    "☐ Consulta Lista 69-B (captura de pantalla)",
                    "☐ Factura CFDI timbrada",
                    "☐ Comprobante de pago",
                ]),
                ("DOCUMENTACIÓN CONTRACTUAL (OBLIGATORIA)", [
                    "☐ Contrato firmado por ambas partes",
                    "☐ Anexo técnico con alcance detallado",
                    "☐ Orden de Compra autorizada",
                ]),
                ("EVIDENCIA DE MATERIALIDAD (OBLIGATORIA)", [
                    "☐ Acta de entrega-recepción firmada",
                    "☐ Entregables físicos/digitales",
                    "☐ Reportes de avance",
                    "☐ Minutas de reuniones",
                    "☐ Correspondencia email del proyecto",
                ]),
                ("EVIDENCIA ADICIONAL (RECOMENDADA)", [
                    "☐ Evidencia fotográfica del trabajo",
                    "☐ Registro de asistencia a reuniones",
                    "☐ Encuesta de satisfacción",
                    "☐ Testimoniales/Referencias",
                ]),
                ("VALIDACIÓN", [
                    "Revisado por: {{REVISOR}}",
                    "Fecha de revisión: {{FECHA_REVISION}}",
                    "Resultado: ☐ Completo ☐ Pendientes",
                ]),
            ]
        },
        {
            "filename": "minutas_trabajo.docx",
            "title": "FORMATO DE MINUTA DE REUNIÓN DE TRABAJO",
            "sections": [
                ("DATOS DE LA REUNIÓN", [
                    f"Fecha: {PLACEHOLDERS['fecha']}",
                    "Hora: {{HORA_REUNION}}",
                    "Lugar/Medio: {{LUGAR_MEDIO}}",
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                ]),
                ("ASISTENTES", [
                    f"Por {PLACEHOLDERS['empresa']}:",
                    "• {{ASISTENTE_CLIENTE_1}} - {{CARGO_1}}",
                    "• {{ASISTENTE_CLIENTE_2}} - {{CARGO_2}}",
                    "",
                    f"Por {PLACEHOLDERS['proveedor']}:",
                    "• {{ASISTENTE_PROVEEDOR_1}} - {{CARGO_3}}",
                    "• {{ASISTENTE_PROVEEDOR_2}} - {{CARGO_4}}",
                ]),
                ("ORDEN DEL DÍA", [
                    "1. {{TEMA_1}}",
                    "2. {{TEMA_2}}",
                    "3. {{TEMA_3}}",
                    "4. Varios y próximos pasos",
                ]),
                ("DESARROLLO DE LA REUNIÓN", [
                    "{{DESARROLLO_REUNION}}",
                ]),
                ("ACUERDOS Y COMPROMISOS", [
                    "| # | Acuerdo | Responsable | Fecha Límite |",
                    "|---|---------|-------------|--------------|",
                    "| 1 | {{ACUERDO_1}} | {{RESP_1}} | {{FECHA_1}} |",
                    "| 2 | {{ACUERDO_2}} | {{RESP_2}} | {{FECHA_2}} |",
                    "| 3 | {{ACUERDO_3}} | {{RESP_3}} | {{FECHA_3}} |",
                ]),
                ("PRÓXIMA REUNIÓN", [
                    "Fecha: {{PROXIMA_FECHA}}",
                    "Temas a tratar: {{PROXIMOS_TEMAS}}",
                ]),
                ("FIRMAS DE CONFORMIDAD", [
                    "",
                    "_________________________          _________________________",
                    f"   {PLACEHOLDERS['empresa']}               {PLACEHOLDERS['proveedor']}",
                ]),
            ]
        },
    ],
    "A7_DEFENSA": [
        {
            "filename": "guia_defense_file.docx",
            "title": "GUÍA DE ESTRUCTURACIÓN DEL DEFENSE FILE",
            "sections": [
                ("OBJETIVO", [
                    "El Defense File es el expediente integral que documenta la materialidad y razón de negocios de una operación, preparado para sustentar la deducibilidad ante una eventual auditoría del SAT.",
                ]),
                ("ESTRUCTURA DEL DEFENSE FILE", [
                    "SECCIÓN 1: INFORMACIÓN GENERAL",
                    "• Carátula con datos del proyecto",
                    "• Índice de documentos",
                    "• Resumen ejecutivo",
                    "",
                    "SECCIÓN 2: DOCUMENTACIÓN FISCAL",
                    "• Análisis de razón de negocios",
                    "• Consulta lista 69-B",
                    "• Opinión de cumplimiento",
                    "",
                    "SECCIÓN 3: DOCUMENTACIÓN CONTRACTUAL",
                    "• Contrato firmado",
                    "• Anexos técnicos",
                    "• Orden de compra",
                    "",
                    "SECCIÓN 4: EVIDENCIA DE MATERIALIDAD",
                    "• Entregables recibidos",
                    "• Actas de entrega",
                    "• Reportes de avance",
                    "• Minutas de reuniones",
                    "",
                    "SECCIÓN 5: DOCUMENTACIÓN FINANCIERA",
                    "• Factura CFDI",
                    "• Comprobante de pago",
                    "• Validación 3-way match",
                ]),
                ("APLICACIÓN", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Empresa: {PLACEHOLDERS['empresa']}",
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                    f"Monto: {PLACEHOLDERS['monto']} MXN",
                ]),
            ]
        },
        {
            "filename": "criterios_tfja.docx",
            "title": "CRITERIOS JURISPRUDENCIALES TFJA APLICABLES",
            "sections": [
                ("TRIBUNAL FEDERAL DE JUSTICIA ADMINISTRATIVA", [
                    "Los siguientes criterios jurisprudenciales son relevantes para la defensa de operaciones de servicios intangibles:",
                ]),
                ("TESIS 1: RAZÓN DE NEGOCIOS", [
                    "Número: VII-TASR-{{NUMERO_TESIS_1}}",
                    "Rubro: RAZÓN DE NEGOCIOS. ELEMENTOS QUE LA INTEGRAN.",
                    "Criterio: La razón de negocios existe cuando el contribuyente demuestra que la operación tiene un propósito comercial legítimo más allá del beneficio fiscal.",
                    "Aplicación: Documentar siempre el beneficio económico esperado de cada operación.",
                ]),
                ("TESIS 2: MATERIALIDAD DE SERVICIOS", [
                    "Número: VII-TASR-{{NUMERO_TESIS_2}}",
                    "Rubro: OPERACIONES INEXISTENTES. CARGA DE LA PRUEBA.",
                    "Criterio: La autoridad debe probar los hechos que sustenten la presunción de inexistencia, mientras que el contribuyente puede desvirtuar la presunción con documentación que acredite la materialidad.",
                    "Aplicación: Mantener evidencia documental robusta de cada servicio prestado.",
                ]),
                ("TESIS 3: ESTRICTA INDISPENSABILIDAD", [
                    "Número: VII-TASR-{{NUMERO_TESIS_3}}",
                    "Rubro: DEDUCCIONES. CONCEPTO DE ESTRICTA INDISPENSABILIDAD.",
                    "Criterio: Un gasto es estrictamente indispensable cuando tiene relación directa con la actividad del contribuyente y es necesario para la consecución de sus fines.",
                    "Aplicación: Documentar la relación entre el servicio contratado y la actividad preponderante.",
                ]),
            ]
        },
        {
            "filename": "documentos_criticos.docx",
            "title": "CHECKLIST DE DOCUMENTOS CRÍTICOS PARA DEFENSA",
            "sections": [
                ("PROYECTO", [
                    f"Nombre: {PLACEHOLDERS['proyecto']}",
                    f"Monto: {PLACEHOLDERS['monto']} MXN",
                    f"Proveedor: {PLACEHOLDERS['proveedor']}",
                ]),
                ("DOCUMENTOS CRÍTICOS - NIVEL 1 (INDISPENSABLES)", [
                    "☐ Contrato de prestación de servicios firmado",
                    "☐ Factura CFDI válida y timbrada",
                    "☐ Comprobante de pago (transferencia bancaria)",
                    "☐ Acta de entrega-recepción firmada",
                    "☐ Consulta lista 69-B del proveedor",
                ]),
                ("DOCUMENTOS CRÍTICOS - NIVEL 2 (ALTAMENTE RECOMENDADOS)", [
                    "☐ Anexo técnico con alcance detallado",
                    "☐ Orden de compra autorizada",
                    "☐ Opinión de cumplimiento fiscal del proveedor",
                    "☐ Análisis de razón de negocios documentado",
                    "☐ Entregables físicos/digitales del servicio",
                ]),
                ("DOCUMENTOS CRÍTICOS - NIVEL 3 (COMPLEMENTARIOS)", [
                    "☐ Reportes de avance del proyecto",
                    "☐ Minutas de reuniones de trabajo",
                    "☐ Correspondencia electrónica del proyecto",
                    "☐ Evidencia fotográfica (si aplica)",
                    "☐ Testimoniales de participantes",
                ]),
                ("VALIDACIÓN DEL EXPEDIENTE", [
                    "Nivel 1 completo: ☐ Sí ☐ No",
                    "Nivel 2 completo: ☐ Sí ☐ No",
                    "Nivel 3 completo: ☐ Sí ☐ No",
                    "",
                    "Fortaleza del expediente:",
                    "☐ ALTA - Niveles 1, 2 y 3 completos",
                    "☐ MEDIA - Niveles 1 y 2 completos",
                    "☐ BAJA - Solo Nivel 1 completo",
                    "☐ CRÍTICA - Nivel 1 incompleto",
                ]),
            ]
        },
        {
            "filename": "precedentes_defensa.docx",
            "title": "PRECEDENTES DE DEFENSA FISCAL EXITOSA",
            "sections": [
                ("CASO 1: CONSULTORÍA ESTRATÉGICA", [
                    "Situación: Auditoría SAT cuestionó deducción de $2.5M en servicios de consultoría.",
                    "Defensa presentada:",
                    "• Contrato detallado con 15 entregables específicos",
                    "• 24 reportes de avance mensuales",
                    "• 12 minutas de reuniones con firmas",
                    "• Análisis de ROI mostrando beneficio de $8M",
                    "Resultado: DEDUCCIÓN CONFIRMADA",
                    "Lección: La documentación exhaustiva de entregables fue determinante.",
                ]),
                ("CASO 2: SERVICIOS DE TECNOLOGÍA", [
                    "Situación: SAT presumió operación simulada por servicios de desarrollo de software.",
                    "Defensa presentada:",
                    "• Repositorio de código con historial de commits",
                    "• Documentación técnica del sistema desarrollado",
                    "• Correos con revisiones y feedback del cliente",
                    "• Testimonio de usuarios del sistema",
                    "Resultado: PRESUNCIÓN DESVIRTUADA",
                    "Lección: La evidencia técnica y testimonial fortaleció la materialidad.",
                ]),
                ("APLICACIÓN AL PROYECTO ACTUAL", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    "Elementos de defensa ya documentados:",
                    "☐ Entregables tangibles",
                    "☐ Comunicaciones del proyecto",
                    "☐ Reportes de avance",
                    "☐ Análisis de beneficio económico",
                ]),
            ]
        },
        {
            "filename": "refuerzo_probatorio.docx",
            "title": "ESTRATEGIAS DE REFUERZO PROBATORIO",
            "sections": [
                ("OBJETIVO", [
                    "Fortalecer el expediente de defensa con elementos adicionales que soporten la materialidad y razón de negocios de las operaciones.",
                ]),
                ("ESTRATEGIA 1: TESTIMONIALES", [
                    "Obtener declaraciones escritas de:",
                    "• Personal del cliente que interactuó con el proveedor",
                    "• Terceros que presenciaron la prestación del servicio",
                    "• Beneficiarios directos del servicio recibido",
                    "",
                    "Formato de testimonial:",
                    "• Nombre completo y cargo del declarante",
                    "• Descripción de su participación en el proyecto",
                    "• Evidencia específica que puede aportar",
                    "• Firma autógrafa y fecha",
                ]),
                ("ESTRATEGIA 2: PERITAJE TÉCNICO", [
                    "Para servicios especializados, considerar:",
                    "• Dictamen pericial sobre la calidad del servicio",
                    "• Valuación independiente del trabajo realizado",
                    "• Certificación de expertos en la materia",
                ]),
                ("ESTRATEGIA 3: DOCUMENTACIÓN COMPLEMENTARIA", [
                    "Elementos adicionales a recopilar:",
                    "☐ Fotografías fechadas del trabajo realizado",
                    "☐ Videos de sesiones de trabajo o capacitación",
                    "☐ Registros de acceso a instalaciones",
                    "☐ Metadatos de archivos entregados",
                    "☐ Historial de comunicaciones (emails, chats)",
                ]),
                ("APLICACIÓN", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    "Estrategias a implementar: {{ESTRATEGIAS_SELECCIONADAS}}",
                    "Responsable de recopilación: {{RESPONSABLE}}",
                    "Fecha límite: {{FECHA_LIMITE}}",
                ]),
            ]
        },
    ],
    "KNOWLEDGE_BASE": [
        {
            "filename": "cff_completo.docx",
            "title": "CÓDIGO FISCAL DE LA FEDERACIÓN - EXTRACTOS RELEVANTES",
            "sections": [
                ("TÍTULO I - DISPOSICIONES GENERALES", [
                    "Artículo 5-A. Los actos jurídicos que carezcan de una razón de negocios y que generen un beneficio fiscal directo o indirecto, tendrán los efectos fiscales que correspondan a los que se habrían realizado para la obtención del beneficio económico razonablemente esperado por el contribuyente.",
                ]),
                ("TÍTULO III - DE LAS FACULTADES DE LAS AUTORIDADES FISCALES", [
                    "Artículo 69-B. Cuando la autoridad fiscal detecte que un contribuyente ha estado emitiendo comprobantes sin contar con los activos, personal, infraestructura o capacidad material, directa o indirectamente, para prestar los servicios o producir, comercializar o entregar los bienes que amparan tales comprobantes, o bien, que dichos contribuyentes se encuentren no localizados, se presumirá la inexistencia de las operaciones amparadas en tales comprobantes.",
                ]),
                ("TÍTULO IV - DE LAS INFRACCIONES Y DELITOS FISCALES", [
                    "Artículo 108. Comete el delito de defraudación fiscal quien con uso de engaños o aprovechamiento de errores, omita total o parcialmente el pago de alguna contribución u obtenga un beneficio indebido con perjuicio del fisco federal.",
                ]),
                ("APLICACIÓN EN " + PLACEHOLDERS['empresa'], [
                    "Los artículos del CFF aplicables a operaciones de servicios intangibles incluyen:",
                    "• Art. 5-A: Razón de negocios",
                    "• Art. 69-B: Materialidad de operaciones",
                    "• Art. 29: Comprobantes fiscales",
                    "• Art. 27: Requisitos de contribuyentes",
                ]),
            ]
        },
        {
            "filename": "lisr_completo.docx",
            "title": "LEY DEL IMPUESTO SOBRE LA RENTA - EXTRACTOS RELEVANTES",
            "sections": [
                ("TÍTULO II - DE LAS PERSONAS MORALES", [
                    "Artículo 25. Los contribuyentes podrán efectuar las siguientes deducciones:",
                    "I. Las devoluciones que se reciban o los descuentos o bonificaciones que se hagan en el ejercicio.",
                    "II. El costo de lo vendido.",
                    "III. Los gastos netos de descuentos, bonificaciones o devoluciones.",
                    "...",
                ]),
                ("REQUISITOS DE DEDUCCIONES", [
                    "Artículo 27. Las deducciones autorizadas en este Título deberán reunir los siguientes requisitos:",
                    "I. Ser estrictamente indispensables para los fines de la actividad del contribuyente.",
                    "II. Estar amparadas con un comprobante fiscal.",
                    "III. Estar debidamente registradas en contabilidad.",
                    "...",
                ]),
                ("GASTOS NO DEDUCIBLES", [
                    "Artículo 28. Para los efectos de este Título, no serán deducibles:",
                    "I. Los pagos por impuesto sobre la renta a cargo del propio contribuyente.",
                    "...",
                    "XXXI. Los pagos que no cumplan con el artículo 27, fracción III del CFF.",
                ]),
            ]
        },
        {
            "filename": "tipologias_servicio.docx",
            "title": "TIPOLOGÍAS DE SERVICIOS INTANGIBLES",
            "sections": [
                ("CATEGORÍA 1: SERVICIOS PROFESIONALES", [
                    "Incluye: Asesoría legal, contable, fiscal, financiera",
                    "Nivel de riesgo: BAJO-MEDIO",
                    "Documentación mínima: Contrato, reportes, factura",
                    "Evidencia típica: Dictámenes, opiniones escritas, análisis",
                ]),
                ("CATEGORÍA 2: CONSULTORÍA ESPECIALIZADA", [
                    "Incluye: Consultoría estratégica, de negocios, organizacional",
                    "Nivel de riesgo: MEDIO",
                    "Documentación mínima: Contrato, plan de trabajo, entregables",
                    "Evidencia típica: Diagnósticos, planes de acción, recomendaciones",
                ]),
                ("CATEGORÍA 3: SERVICIOS TECNOLÓGICOS", [
                    "Incluye: Desarrollo de software, implementación de sistemas, TI",
                    "Nivel de riesgo: BAJO",
                    "Documentación mínima: Contrato, especificaciones técnicas, código fuente",
                    "Evidencia típica: Sistemas funcionando, documentación técnica, código",
                ]),
                ("CATEGORÍA 4: SERVICIOS DE INTERMEDIACIÓN", [
                    "Incluye: Comisiones, gestión comercial, representación",
                    "Nivel de riesgo: ALTO",
                    "Documentación mínima: Contrato, reportes de gestión, resultados",
                    "Evidencia típica: Contratos conseguidos, ventas generadas, minutas",
                ]),
                ("CATEGORÍA 5: SERVICIOS DE CAPACITACIÓN", [
                    "Incluye: Cursos, talleres, coaching, formación",
                    "Nivel de riesgo: MEDIO",
                    "Documentación mínima: Programa, listas de asistencia, evaluaciones",
                    "Evidencia típica: Material didáctico, fotos, certificados",
                ]),
            ]
        },
        {
            "filename": "umbrales_riesgo.docx",
            "title": "MATRIZ DE UMBRALES DE RIESGO FISCAL",
            "sections": [
                ("UMBRALES POR MONTO", [
                    "BAJO (Verde): Hasta $250,000 MXN",
                    "• Validación básica de proveedor",
                    "• Contrato simple",
                    "",
                    "MEDIO (Amarillo): $250,001 - $1,000,000 MXN",
                    "• Validación completa 69-B",
                    "• Contrato detallado con anexos",
                    "• Análisis de razón de negocios",
                    "",
                    "ALTO (Naranja): $1,000,001 - $5,000,000 MXN",
                    "• Defense File completo obligatorio",
                    "• Dictámenes de todos los agentes",
                    "• Revisión legal exhaustiva",
                    "",
                    "CRÍTICO (Rojo): Más de $5,000,000 MXN",
                    "• Aprobación de Consejo",
                    "• Defense File con validación externa",
                    "• Seguimiento post-pago por 5 años",
                ]),
                ("FACTORES AGRAVANTES", [
                    "Cada factor suma +1 nivel de riesgo:",
                    "☐ Proveedor de reciente creación (<2 años)",
                    "☐ Servicio 100% intangible sin entregables claros",
                    "☐ Operación con partes relacionadas",
                    "☐ Proveedor con domicilio no verificable",
                    "☐ Primera operación con el proveedor",
                    "☐ Monto atípico vs histórico",
                ]),
                ("EVALUACIÓN DEL PROYECTO", [
                    f"Proyecto: {PLACEHOLDERS['proyecto']}",
                    f"Monto: {PLACEHOLDERS['monto']} MXN",
                    "Nivel base por monto: {{NIVEL_BASE}}",
                    "Factores agravantes: {{FACTORES_AGRAVANTES}}",
                    "NIVEL DE RIESGO FINAL: {{NIVEL_RIESGO_FINAL}}",
                ]),
            ]
        },
        {
            "filename": "politica_anti_efos.docx",
            "title": "POLÍTICA ANTI-EFOS DE LA EMPRESA",
            "sections": [
                ("OBJETIVO", [
                    f"Establecer los lineamientos que {PLACEHOLDERS['empresa']} seguirá para prevenir la contratación de Empresas que Facturan Operaciones Simuladas (EFOS).",
                ]),
                ("DEFINICIÓN", [
                    "EFOS: Contribuyentes que emiten comprobantes fiscales sin tener capacidad material para prestar los servicios o entregar los bienes que amparan dichos comprobantes.",
                ]),
                ("PROCEDIMIENTO DE VALIDACIÓN", [
                    "ANTES de cualquier contratación de servicios:",
                    "1. Consultar lista 69-B del SAT",
                    "2. Verificar opinión de cumplimiento fiscal",
                    "3. Validar domicilio fiscal del proveedor",
                    "4. Confirmar existencia de personal y activos",
                    "5. Revisar antigüedad y objeto social del RFC",
                ]),
                ("PROHIBICIONES", [
                    "Está prohibido contratar proveedores que:",
                    "☐ Aparezcan en lista definitiva 69-B",
                    "☐ No proporcionen opinión de cumplimiento positiva",
                    "☐ No puedan demostrar capacidad operativa",
                    "☐ Tengan domicilio fiscal no localizable",
                    "☐ Tengan RFC con antigüedad menor a 1 año para montos >$500,000",
                ]),
                ("SANCIONES INTERNAS", [
                    "El incumplimiento de esta política resultará en:",
                    "• Primera falta: Amonestación escrita",
                    "• Segunda falta: Suspensión temporal",
                    "• Tercera falta: Rescisión de contrato laboral",
                    "• En todos los casos: Responsabilidad solidaria por contingencias fiscales",
                ]),
                ("VIGENCIA", [
                    f"Esta política entra en vigor a partir de: {PLACEHOLDERS['fecha']}",
                    f"Aprobada por: {{DIRECTOR_GENERAL}}",
                ]),
            ]
        },
    ],
}


def generate_document(agent_id: str, template_config: dict) -> str:
    """Generate a single document"""
    doc = Document()
    
    create_header(doc, template_config["title"], f"Empresa: {PLACEHOLDERS['empresa']}")
    
    for section_title, content in template_config["sections"]:
        add_section(doc, section_title, content)
    
    create_footer(doc)
    
    agent_dir = TEMPLATES_DIR / agent_id
    agent_dir.mkdir(parents=True, exist_ok=True)
    
    file_path = agent_dir / template_config["filename"]
    doc.save(str(file_path))
    
    return str(file_path)


def generate_manifest(agent_id: str, templates: list) -> str:
    """Generate manifest.json for an agent"""
    manifest = {
        "agent_id": agent_id,
        "version": "1.0.0",
        "generated_at": "{{FECHA_ACTUAL}}",
        "templates": []
    }
    
    for template in templates:
        manifest["templates"].append({
            "filename": template["filename"],
            "title": template["title"],
            "section_count": len(template["sections"]),
            "placeholders": list(set(
                placeholder 
                for _, content in template["sections"] 
                for line in content 
                if "{{" in line 
                for placeholder in [p.split("}}")[0] for p in line.split("{{")[1:]]
            ))
        })
    
    agent_dir = TEMPLATES_DIR / agent_id
    manifest_path = agent_dir / "manifest.json"
    
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    
    return str(manifest_path)


def main():
    """Generate all templates"""
    print("=" * 60)
    print("GENERANDO 42 TEMPLATES PARA REVISAR.IA")
    print("=" * 60)
    
    total_generated = 0
    
    for agent_id, templates in TEMPLATES.items():
        print(f"\n[{agent_id}] Generando {len(templates)} templates...")
        
        for template in templates:
            try:
                path = generate_document(agent_id, template)
                print(f"  ✓ {template['filename']}")
                total_generated += 1
            except Exception as e:
                print(f"  ✗ {template['filename']}: {e}")
        
        manifest_path = generate_manifest(agent_id, templates)
        print(f"  ✓ manifest.json")
    
    print("\n" + "=" * 60)
    print(f"COMPLETADO: {total_generated} templates generados")
    print("=" * 60)


if __name__ == "__main__":
    main()
