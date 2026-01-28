"""
Ingestion Service (A-ING)
Text extraction and processing for knowledge repository documents
"""
import os
import re
import uuid
import json
import logging
from datetime import datetime
from typing import Optional, Dict, Any, Tuple
import asyncpg

logger = logging.getLogger(__name__)

DATABASE_URL = os.environ.get('DATABASE_URL', '')


def safe_uuid(id_string: str) -> uuid.UUID:
    """Convert ID string to UUID safely, generating deterministic UUID for non-UUID strings."""
    if not id_string:
        raise ValueError("ID string cannot be empty")
    try:
        return uuid.UUID(id_string)
    except (ValueError, AttributeError):
        namespace = uuid.UUID('6ba7b810-9dad-11d1-80b4-00c04fd430c8')
        return uuid.uuid5(namespace, id_string)


async def get_db_connection():
    """Get asyncpg connection for database operations"""
    if not DATABASE_URL:
        logger.warning("DATABASE_URL not configured")
        return None
    try:
        return await asyncpg.connect(DATABASE_URL)
    except Exception as e:
        logger.error(f"Database connection failed: {e}")
        return None


class IngestionService:
    """Service for extracting and processing text from uploaded documents."""
    
    def __init__(self):
        self.supported_extensions = {
            '.pdf': self.extract_pdf,
            '.docx': self.extract_docx,
            '.xlsx': self.extract_xlsx,
            '.xls': self.extract_xlsx,
            '.txt': self.extract_txt,
            '.md': self.extract_txt,
            '.csv': self.extract_txt,
        }
    
    async def process_document(
        self, 
        document_id: str, 
        empresa_id: str, 
        file_path: str
    ) -> Dict[str, Any]:
        """
        Main entry point for document processing.
        Extracts text, normalizes it, and stores in database.
        """
        logger.info(f"Starting ingestion for document {document_id}")
        
        try:
            await self._update_document_status(document_id, empresa_id, 'processing')
            
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            _, ext = os.path.splitext(file_path.lower())
            
            if ext not in self.supported_extensions:
                logger.warning(f"Unsupported file extension: {ext}")
                await self._update_document_status(document_id, empresa_id, 'uploaded')
                return {
                    "success": False,
                    "message": f"Unsupported file type: {ext}",
                    "document_id": document_id
                }
            
            extract_func = self.supported_extensions[ext]
            extracted_text, extraction_meta = await extract_func(file_path)
            
            normalized_text = self.normalize_text(extracted_text)
            
            word_count = len(normalized_text.split()) if normalized_text else 0
            page_count = extraction_meta.get('page_count', 1)
            extraction_method = extraction_meta.get('method', 'unknown')
            
            await self.store_extracted_text(
                document_id=document_id,
                empresa_id=empresa_id,
                text=normalized_text,
                metadata={
                    "language": "es",
                    "page_count": page_count,
                    "word_count": word_count,
                    "extraction_method": extraction_method,
                    "extraction_meta": extraction_meta
                }
            )
            
            await self._update_document_status(document_id, empresa_id, 'indexed')
            
            logger.info(f"Successfully processed document {document_id}: {word_count} words extracted")
            
            return {
                "success": True,
                "document_id": document_id,
                "word_count": word_count,
                "page_count": page_count,
                "extraction_method": extraction_method
            }
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            await self._update_document_status(document_id, empresa_id, 'error')
            return {
                "success": False,
                "document_id": document_id,
                "error": str(e)
            }
    
    async def extract_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using pdfplumber, with OCR fallback."""
        import pdfplumber
        
        text_parts = []
        page_count = 0
        method = 'pdfplumber'
        
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
            
            full_text = "\n\n".join(text_parts)
            
            if len(full_text.strip()) < 100 and page_count > 0:
                logger.info(f"PDF has little text, attempting OCR fallback for {file_path}")
                ocr_text, ocr_meta = await self._extract_pdf_ocr(file_path)
                if len(ocr_text) > len(full_text):
                    full_text = ocr_text
                    method = 'ocr'
                    page_count = ocr_meta.get('page_count', page_count)
            
            return full_text, {
                "method": method,
                "page_count": page_count,
                "file_path": file_path
            }
            
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying OCR: {e}")
            return await self._extract_pdf_ocr(file_path)
    
    async def _extract_pdf_ocr(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using OCR (pytesseract)."""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            images = convert_from_path(file_path, dpi=200)
            page_count = len(images)
            
            text_parts = []
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image, lang='spa+eng')
                text_parts.append(page_text)
            
            return "\n\n".join(text_parts), {
                "method": "ocr",
                "page_count": page_count,
                "file_path": file_path,
                "ocr_lang": "spa+eng"
            }
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return "", {
                "method": "ocr_failed",
                "page_count": 0,
                "error": str(e)
            }
    
    async def extract_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX using python-docx."""
        from docx import Document
        
        try:
            doc = Document(file_path)
            
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            
            return full_text, {
                "method": "python-docx",
                "page_count": 1,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return "", {
                "method": "python-docx",
                "error": str(e)
            }
    
    async def extract_xlsx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Excel files using openpyxl."""
        from openpyxl import load_workbook
        
        try:
            wb = load_workbook(file_path, data_only=True)
            
            text_parts = []
            sheet_count = 0
            row_count = 0
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_count += 1
                text_parts.append(f"=== Hoja: {sheet_name} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(v.strip() for v in row_values):
                        text_parts.append(" | ".join(row_values))
                        row_count += 1
            
            full_text = "\n".join(text_parts)
            
            return full_text, {
                "method": "openpyxl",
                "page_count": sheet_count,
                "sheet_count": sheet_count,
                "row_count": row_count
            }
            
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            return "", {
                "method": "openpyxl",
                "error": str(e)
            }
    
    async def extract_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text files."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                return content, {
                    "method": "plain_text",
                    "page_count": 1,
                    "encoding": encoding
                }
                
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Text extraction failed: {e}")
                return "", {
                    "method": "plain_text",
                    "error": str(e)
                }
        
        return "", {
            "method": "plain_text",
            "error": "Could not decode file with any supported encoding"
        }
    
    def normalize_text(self, text: str) -> str:
        """Clean whitespace and fix encoding issues."""
        if not text:
            return ""
        
        text = text.replace('\x00', '')
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        text = re.sub(r'[ \t]+', ' ', text)
        
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines]
        text = '\n'.join(cleaned_lines)
        
        text = text.strip()
        
        return text
    
    async def store_extracted_text(
        self,
        document_id: str,
        empresa_id: str,
        text: str,
        metadata: Dict[str, Any]
    ) -> bool:
        """Save extracted text to the database."""
        conn = await get_db_connection()
        if not conn:
            raise Exception("Database connection unavailable")
        
        try:
            existing = await conn.fetchrow(
                """
                SELECT id FROM knowledge_document_text
                WHERE document_id = $1 AND empresa_id = $2
                """,
                uuid.UUID(document_id),
                safe_uuid(empresa_id)
            )
            
            if existing:
                await conn.execute(
                    """
                    UPDATE knowledge_document_text
                    SET extracted_text = $1,
                        language = $2,
                        page_count = $3,
                        word_count = $4,
                        extraction_method = $5,
                        extraction_meta = $6,
                        updated_at = NOW()
                    WHERE document_id = $7 AND empresa_id = $8
                    """,
                    text,
                    metadata.get('language', 'es'),
                    metadata.get('page_count'),
                    metadata.get('word_count'),
                    metadata.get('extraction_method'),
                    json.dumps(metadata.get('extraction_meta', {})),
                    uuid.UUID(document_id),
                    safe_uuid(empresa_id)
                )
            else:
                text_id = uuid.uuid4()
                await conn.execute(
                    """
                    INSERT INTO knowledge_document_text
                    (id, document_id, empresa_id, extracted_text, language, page_count, word_count, extraction_method, extraction_meta, created_at, updated_at)
                    VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, NOW(), NOW())
                    """,
                    text_id,
                    uuid.UUID(document_id),
                    safe_uuid(empresa_id),
                    text,
                    metadata.get('language', 'es'),
                    metadata.get('page_count'),
                    metadata.get('word_count'),
                    metadata.get('extraction_method'),
                    json.dumps(metadata.get('extraction_meta', {}))
                )
            
            logger.info(f"Stored extracted text for document {document_id}")
            return True
            
        finally:
            await conn.close()
    
    async def _update_document_status(
        self, 
        document_id: str, 
        empresa_id: str, 
        status: str
    ) -> bool:
        """Update document processing status."""
        conn = await get_db_connection()
        if not conn:
            logger.warning("Cannot update status - database unavailable")
            return False
        
        try:
            await conn.execute(
                """
                UPDATE knowledge_documents 
                SET status = $1, updated_at = NOW()
                WHERE id = $2 AND empresa_id = $3
                """,
                status,
                uuid.UUID(document_id),
                safe_uuid(empresa_id)
            )
            logger.debug(f"Updated document {document_id} status to {status}")
            return True
        except Exception as e:
            logger.error(f"Failed to update document status: {e}")
            return False
        finally:
            await conn.close()
    
    async def get_extracted_text(
        self, 
        document_id: str, 
        empresa_id: str
    ) -> Optional[Dict[str, Any]]:
        """Retrieve extracted text for a document."""
        conn = await get_db_connection()
        if not conn:
            raise Exception("Database connection unavailable")
        
        try:
            row = await conn.fetchrow(
                """
                SELECT id, document_id, empresa_id, extracted_text, language,
                       page_count, word_count, extraction_method, extraction_meta,
                       created_at, updated_at
                FROM knowledge_document_text
                WHERE document_id = $1 AND empresa_id = $2
                """,
                uuid.UUID(document_id),
                safe_uuid(empresa_id)
            )
            
            if not row:
                return None
            
            return {
                "id": str(row['id']),
                "document_id": str(row['document_id']),
                "empresa_id": str(row['empresa_id']),
                "extracted_text": row['extracted_text'],
                "language": row['language'],
                "page_count": row['page_count'],
                "word_count": row['word_count'],
                "extraction_method": row['extraction_method'],
                "extraction_meta": row['extraction_meta'] if row['extraction_meta'] else {},
                "created_at": row['created_at'].isoformat() if row['created_at'] else None,
                "updated_at": row['updated_at'].isoformat() if row['updated_at'] else None
            }
            
        finally:
            await conn.close()


ingestion_service = IngestionService()
