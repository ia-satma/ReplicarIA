"""
Template Engine for RAG Document Generation
Replaces {{PLACEHOLDER}} with actual data in .docx templates
"""
import re
import json
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from copy import deepcopy
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

PLACEHOLDER_PATTERN = re.compile(r'\{\{([A-Z_0-9]+)\}\}')


def extract_placeholders(text: str) -> List[str]:
    """Extract all placeholder names from text"""
    return PLACEHOLDER_PATTERN.findall(text)


def replace_placeholders(text: str, data: Dict[str, str]) -> str:
    """Replace all {{PLACEHOLDER}} with corresponding data values"""
    def replacer(match):
        key = match.group(1)
        return str(data.get(key, match.group(0)))
    
    return PLACEHOLDER_PATTERN.sub(replacer, text)


def process_paragraph(paragraph: Paragraph, data: Dict[str, str]) -> None:
    """Process a paragraph and replace placeholders in all runs"""
    full_text = paragraph.text
    if not PLACEHOLDER_PATTERN.search(full_text):
        return
    
    new_text = replace_placeholders(full_text, data)
    
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def process_table(table: Table, data: Dict[str, str]) -> None:
    """Process all cells in a table"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph, data)
            for nested_table in cell.tables:
                process_table(nested_table, data)


def process_header_footer(section, data: Dict[str, str]) -> None:
    """Process headers and footers in a section"""
    if section.header:
        for paragraph in section.header.paragraphs:
            process_paragraph(paragraph, data)
        for table in section.header.tables:
            process_table(table, data)
    
    if section.footer:
        for paragraph in section.footer.paragraphs:
            process_paragraph(paragraph, data)
        for table in section.footer.tables:
            process_table(table, data)


def get_template_placeholders(agent_id: str, template_id: str) -> List[str]:
    """Get all placeholders from a template by reading its manifest or scanning the document"""
    agent_dir = TEMPLATES_DIR / agent_id
    manifest_path = agent_dir / "manifest.json"
    
    if manifest_path.exists():
        try:
            with open(manifest_path, "r", encoding="utf-8") as f:
                manifest = json.load(f)
                for t in manifest.get("templates", []):
                    if t.get("filename") == f"{template_id}.docx":
                        return t.get("placeholders", [])
        except (json.JSONDecodeError, IOError):
            pass
    
    template_path = agent_dir / f"{template_id}.docx"
    if template_path.exists():
        return scan_document_placeholders(template_path)
    
    return []


def scan_document_placeholders(template_path: Path) -> List[str]:
    """Scan a document to find all unique placeholders"""
    placeholders = set()
    
    try:
        doc = Document(str(template_path))
        
        for paragraph in doc.paragraphs:
            placeholders.update(extract_placeholders(paragraph.text))
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(extract_placeholders(paragraph.text))
        
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    placeholders.update(extract_placeholders(paragraph.text))
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    placeholders.update(extract_placeholders(paragraph.text))
    
    except Exception as e:
        logger.error(f"Error scanning document {template_path}: {e}")
    
    return list(placeholders)


def validate_data(agent_id: str, template_id: str, data: Dict[str, str]) -> Dict[str, Any]:
    """
    Validate provided data against template placeholders.
    Returns dict with:
    - valid: bool
    - missing: list of missing required fields
    - extra: list of provided fields not in template
    - placeholders: list of all template placeholders
    """
    template_placeholders = get_template_placeholders(agent_id, template_id)
    
    provided_keys = set(data.keys())
    template_keys = set(template_placeholders)
    
    missing = list(template_keys - provided_keys)
    extra = list(provided_keys - template_keys)
    
    return {
        "valid": len(missing) == 0,
        "missing": sorted(missing),
        "extra": sorted(extra),
        "placeholders": sorted(template_placeholders),
        "provided": sorted(list(provided_keys))
    }


def fill_template(template_path: Path, data: Dict[str, str]) -> Path:
    """
    Fill a template with provided data and return path to generated document.
    
    Args:
        template_path: Path to the .docx template
        data: Dictionary mapping placeholder names to values
        
    Returns:
        Path to the generated document (in temp directory)
    """
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    doc = Document(str(template_path))
    
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, data)
    
    for table in doc.tables:
        process_table(table, data)
    
    for section in doc.sections:
        process_header_footer(section, data)
    
    temp_dir = Path(tempfile.gettempdir()) / "revisar_ia_generated"
    temp_dir.mkdir(exist_ok=True)
    
    output_filename = f"{template_path.stem}_filled.docx"
    output_path = temp_dir / output_filename
    
    doc.save(str(output_path))
    
    return output_path


def generate_document(
    agent_id: str, 
    template_id: str, 
    data: Dict[str, str],
    validate: bool = True
) -> Tuple[Path, Dict[str, Any]]:
    """
    Generate a filled document from a template.
    
    Args:
        agent_id: Agent folder name (e.g., "A3_FISCAL")
        template_id: Template file name without extension
        data: Dictionary with placeholder values
        validate: Whether to validate data before filling
        
    Returns:
        Tuple of (output_path, validation_result)
    """
    agent_dir = TEMPLATES_DIR / agent_id
    template_path = agent_dir / f"{template_id}.docx"
    
    if not agent_dir.exists():
        raise ValueError(f"Agent not found: {agent_id}")
    
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_id}")
    
    validation = validate_data(agent_id, template_id, data) if validate else {"valid": True}
    
    output_path = fill_template(template_path, data)
    
    return output_path, validation
