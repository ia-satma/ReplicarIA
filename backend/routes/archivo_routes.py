from fastapi import APIRouter, HTTPException, UploadFile, File, Form, Depends
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, EmailStr, field_validator
from typing import List, Optional
import json
import logging
import uuid
import os
import re
import io
import tempfile
from datetime import datetime, timezone

from jose import jwt, exceptions as jose_exceptions
from services.archivo_chat_service import archivo_service
from services.user_db import user_service

try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pytesseract
    from pdf2image import convert_from_bytes
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from services.auth_service import get_secret_key, security

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/chat", tags=["archivo"])
onboarding_router = APIRouter(prefix="/api/onboarding", tags=["onboarding"])

# Usar servicio centralizado de autenticación
SECRET_KEY = get_secret_key()
ALGORITHM = "HS256"


class ChatMessage(BaseModel):
    message: str
    history: Optional[List[dict]] = []
    company_id: Optional[str] = None
    company_context: Optional[dict] = None
    user_context: Optional[dict] = None


class DocumentAnalysisResponse(BaseModel):
    success: bool
    classification: str
    file_name: str
    summary: Optional[str] = None
    extracted_data: Optional[dict] = None
    error: Optional[str] = None


async def get_user_from_token(credentials: HTTPAuthorizationCredentials = Depends(security)) -> Optional[dict]:
    """Obtiene usuario actual del token JWT o sesión OTP - retorna None si no hay token"""
    if not credentials:
        return None
    
    token = credentials.credentials
    
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub") or payload.get("user_id")
        if user_id:
            user = await user_service.get_user_by_id(user_id)
            if user:
                user_dict = user.to_dict() if hasattr(user, 'to_dict') else dict(user)
                return {
                    'id': user_dict.get('id'),
                    'email': user_dict.get('email'),
                    'nombre': user_dict.get('nombre') or user_dict.get('full_name') or 'Usuario',
                    'is_admin': user_dict.get('is_admin', False),
                    'cliente_id': user_dict.get('cliente_id'),
                    'cliente_nombre': user_dict.get('cliente_nombre')
                }
    except Exception as e:
        logger.debug(f"Error decodificando token: {e}")
    
    try:
        from services.otp_auth_service import otp_auth_service
        session = await otp_auth_service.validate_session(token)
        if session and session.get("user"):
            u = session["user"]
            return {
                'id': u.get('id'),
                'email': u.get('email'),
                'nombre': u.get('nombre') or 'Usuario',
                'is_admin': u.get('is_admin', False),
                'cliente_id': u.get('cliente_id'),
                'cliente_nombre': u.get('cliente_nombre')
            }
    except Exception as e:
        logger.debug(f"Error validando sesión OTP: {e}")
    
    return None


@router.post("/archivo")
async def chat_archivo(
    data: ChatMessage,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """
    Endpoint principal de chat con ARCHIVO
    Soporta streaming de respuestas y contexto de usuario
    """
    user_context = await get_user_from_token(credentials)
    
    if data.user_context:
        if user_context:
            user_context.update(data.user_context)
        else:
            user_context = data.user_context
    
    async def generate():
        try:
            async for chunk in archivo_service.chat_stream(
                message=data.message,
                history=data.history,
                company_id=data.company_id,
                company_context=data.company_context,
                user_context=user_context
            ):
                yield chunk
        except Exception as e:
            logger.error(f"Chat streaming error: {e}")
            yield f"⚠️ Error: {str(e)}"
    
    return StreamingResponse(
        generate(), 
        media_type="text/plain",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no"
        }
    )


MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB limit

ALLOWED_MIME_TYPES = {
    'application/pdf': 'pdf',
    'image/png': 'image',
    'image/jpeg': 'image',
    'image/jpg': 'image',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/msword': 'doc',
    'application/octet-stream': 'auto',  # Will detect by extension
    'application/x-zip-compressed': 'docx',  # Some browsers send this for DOCX
    'application/zip': 'docx',  # DOCX is a ZIP file
}

# Extension to file type mapping (fallback when MIME is generic)
EXTENSION_TO_TYPE = {
    '.pdf': 'pdf',
    '.png': 'image',
    '.jpg': 'image',
    '.jpeg': 'image',
    '.docx': 'docx',
    '.doc': 'doc',
    '.xlsx': 'excel',
    '.xls': 'excel',
    '.txt': 'text',
}


def detect_file_type(filename: str, content_type: str) -> tuple[str, bool]:
    """
    Detect file type from MIME type or extension.
    Returns (file_type, is_valid)
    """
    file_ext = '.' + filename.split('.')[-1].lower() if '.' in filename else ''

    # Try MIME type first
    file_type = ALLOWED_MIME_TYPES.get(content_type)

    # If MIME is generic or unknown, detect by extension
    if file_type == 'auto' or file_type is None:
        file_type = EXTENSION_TO_TYPE.get(file_ext)
        if file_type:
            logger.info(f"Detected file type '{file_type}' from extension '{file_ext}' (MIME was {content_type})")

    return file_type, file_type is not None

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX extraction disabled")


def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, str]:
    """
    Extract text from PDF using PyMuPDF first, then OCR as fallback.
    Returns (text, extraction_method)
    """
    text = ""
    method = "none"
    
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                text += page.get_text()
            doc.close()
            
            if text.strip():
                method = "pymupdf"
                return text.strip(), method
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
    
    if OCR_AVAILABLE and not text.strip():
        try:
            images = convert_from_bytes(file_bytes)
            ocr_text_parts = []
            for img in images:
                ocr_text_parts.append(pytesseract.image_to_string(img, lang='spa+eng'))
            text = "\n\n".join(ocr_text_parts)
            if text.strip():
                method = "ocr"
        except Exception as e:
            logger.warning(f"OCR extraction failed: {e}")
    
    return text.strip(), method


def extract_text_from_image(file_bytes: bytes) -> tuple[str, str]:
    """
    Extract text from image using pytesseract OCR.
    Returns (text, extraction_method)
    """
    if not OCR_AVAILABLE:
        return "", "none"
    
    try:
        img = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(img, lang='spa+eng')
        return text.strip(), "ocr"
    except Exception as e:
        logger.warning(f"Image OCR failed: {e}")
        return "", "none"


def extract_text_from_docx(file_bytes: bytes) -> tuple[str, str]:
    """
    Extract text from DOCX using multiple fallback methods.
    Returns (text, extraction_method)
    """
    logger.info(f"DOCX extraction starting, content size: {len(file_bytes)} bytes")

    # Method 1: Try with python-docx via BytesIO
    if DOCX_AVAILABLE:
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            text = "\n".join(text_parts)
            if text.strip():
                logger.info(f"DOCX BytesIO extracted: {len(text)} chars")
                return text.strip()[:30000], "python-docx-bytesio"
        except Exception as e:
            logger.warning(f"DOCX BytesIO extraction failed: {e}")

    # Method 2: Try with temp file
    if DOCX_AVAILABLE:
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            doc = DocxDocument(tmp_path)
            text_parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            text = "\n".join(text_parts)
            os.unlink(tmp_path)
            if text.strip():
                logger.info(f"DOCX temp file extracted: {len(text)} chars")
                return text.strip()[:30000], "python-docx-tempfile"
        except Exception as e2:
            logger.warning(f"DOCX temp file extraction failed: {e2}")
            try:
                os.unlink(tmp_path)
            except (OSError, NameError):
                pass

    # Method 3: Raw XML extraction from ZIP (DOCX is a ZIP with XML files)
    try:
        import zipfile
        import re as regex

        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            if 'word/document.xml' in zf.namelist():
                xml_content = zf.read('word/document.xml').decode('utf-8', errors='ignore')
                text_matches = regex.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                text = ' '.join(text_matches)
                text = regex.sub(r'\s+', ' ', text).strip()
                if text:
                    logger.info(f"DOCX raw XML extracted: {len(text)} chars")
                    return text[:30000], "xml-extraction"
    except Exception as e3:
        logger.warning(f"DOCX raw XML extraction failed: {e3}")

    # Method 4: Try to decode as plain text (in case it's not a real DOCX)
    try:
        text = file_bytes.decode('utf-8', errors='ignore')[:30000]
        if text and len([c for c in text[:1000] if c.isalpha()]) > 100:
            logger.info(f"DOCX decoded as plain text: {len(text)} chars")
            return text, "plain-text"
    except (UnicodeDecodeError, AttributeError):
        pass

    logger.error("All DOCX extraction methods failed")
    return "", "none"


@router.post("/archivo/analyze-document")
async def analyze_document(
    file: UploadFile = File(...),
):
    """
    Analiza un documento subido:
    1. Valida tipo y tamaño del archivo
    2. Extrae texto usando PyMuPDF o OCR
    3. Clasifica el documento
    4. Retorna texto extraído, clasificación y resumen
    """
    try:
        filename = file.filename or "unknown"
        file_ext = '.' + filename.split('.')[-1].lower() if '.' in filename else ''

        # Determine file type from MIME or extension
        file_type = ALLOWED_MIME_TYPES.get(file.content_type)

        # If MIME is generic (octet-stream) or unknown, detect by extension
        if file_type == 'auto' or file_type is None:
            file_type = EXTENSION_TO_TYPE.get(file_ext)
            if file_type:
                logger.info(f"Detected file type '{file_type}' from extension '{file_ext}' (MIME was {file.content_type})")

        if not file_type:
            raise HTTPException(
                status_code=400,
                detail=f"Tipo de archivo no permitido: {file.content_type} / {file_ext}. Solo se aceptan PDF, DOCX, DOC, PNG, JPG, JPEG, TXT."
            )

        file_bytes = await file.read()

        if len(file_bytes) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"El archivo excede el límite de 50MB. Tamaño: {len(file_bytes) / (1024*1024):.2f}MB"
            )

        extracted_text = ""
        extraction_method = "none"

        if file_type == 'pdf':
            extracted_text, extraction_method = extract_text_from_pdf(file_bytes)
        elif file_type == 'image':
            extracted_text, extraction_method = extract_text_from_image(file_bytes)
        elif file_type == 'text':
            extracted_text = file_bytes.decode('utf-8', errors='ignore')[:30000]
            extraction_method = "plain-text"
        elif file_type in ['docx', 'doc']:
            extracted_text, extraction_method = extract_text_from_docx(file_bytes)
        
        classification_result = archivo_service.classify_document_by_name(
            file_name=file.filename,
            file_type=file.content_type
        )
        
        text_preview = extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text

        word_count = len(extracted_text.split()) if extracted_text else 0

        # Extract structured data using deep_research_service
        extracted_data = {}
        field_confidence = {}
        investigation_performed = False
        investigation_fields = []

        if extracted_text and len(extracted_text) > 50:
            try:
                from services.deep_research_service import deep_research_service
                classification_type = classification_result.get("classification", "otro")
                doc_analysis = await deep_research_service.analizar_documento(
                    contenido=extracted_text,
                    tipo=classification_type
                )
                if doc_analysis.get("success"):
                    extracted_data = doc_analysis.get("datos", {})
                    field_confidence = doc_analysis.get("confianza_campos", {})
                    logger.info(f"Extracted structured data: {list(extracted_data.keys())}")
            except Exception as e:
                logger.warning(f"Could not extract structured data: {e}")

        # Auto-trigger Deep Research if critical fields are missing
        campos_criticos_faltantes = []
        if not extracted_data.get("rfc"):
            campos_criticos_faltantes.append("rfc")
        if not extracted_data.get("nombre") and not extracted_data.get("razon_social"):
            campos_criticos_faltantes.append("nombre")

        # Only do deep research if we have some starting data to work with
        tiene_datos_base = extracted_data.get("nombre") or extracted_data.get("razon_social") or extracted_data.get("rfc") or extracted_data.get("email")

        if campos_criticos_faltantes and tiene_datos_base:
            try:
                from services.deep_research_service import deep_research_service
                logger.info(f"Auto-triggering Deep Research for missing fields: {campos_criticos_faltantes}")

                investigacion = await deep_research_service.investigar_empresa(
                    nombre=extracted_data.get("nombre") or extracted_data.get("razon_social"),
                    rfc=extracted_data.get("rfc"),
                    documento=extracted_text[:5000]  # First 5000 chars for context
                )

                if investigacion.get("data"):
                    datos_investigacion = investigacion["data"]
                    investigation_performed = True

                    # Merge found data into extracted_data
                    for campo, valor in datos_investigacion.items():
                        if valor and not extracted_data.get(campo):
                            extracted_data[campo] = valor
                            investigation_fields.append(campo)
                            # Add confidence from investigation
                            conf = investigacion.get("field_confidence", {}).get(campo, {})
                            if conf:
                                field_confidence[campo] = conf.get("confidence", 70)
                            logger.info(f"Deep Research found: {campo} = {valor}")

            except Exception as e:
                logger.warning(f"Deep Research failed (non-critical): {e}")

        return {
            "success": True,
            "classification": classification_result.get("classification", "otro"),
            "file_name": file.filename,
            "file_type": file.content_type,
            "file_size_mb": round(len(file_bytes) / (1024 * 1024), 2),
            "summary": classification_result.get("summary", ""),
            "extracted_text": extracted_text,
            "text_preview": text_preview,
            "extraction_method": extraction_method,
            "word_count": word_count,
            "ocr_available": OCR_AVAILABLE,
            "pymupdf_available": PYMUPDF_AVAILABLE,
            "docx_available": DOCX_AVAILABLE,
            "extracted_data": extracted_data,
            "field_confidence": field_confidence,
            "deep_research_performed": investigation_performed,
            "deep_research_fields": investigation_fields,
            "missing_critical_fields": [f for f in campos_criticos_faltantes if f not in extracted_data]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document analysis error: {e}")
        # Determine appropriate status code based on error type
        if "file" in str(e).lower() or "not found" in str(e).lower():
            status_code = 404
        elif "timeout" in str(e).lower() or "memory" in str(e).lower():
            status_code = 503
        else:
            status_code = 400
        raise HTTPException(
            status_code=status_code,
            detail=f"Error analizando documento: {str(e)}"
        )


@router.post("/archivo/extract-data")
async def extract_company_data(data: ChatMessage):
    """
    Extrae datos estructurados de la empresa del historial de conversación
    """
    try:
        extracted = archivo_service.extract_company_data(data.history or [])
        return {
            "success": True,
            "data": extracted
        }
    except Exception as e:
        logger.error(f"Data extraction error: {e}")
        return {
            "success": False,
            "error": str(e),
            "data": {}
        }


class OnboardingData(BaseModel):
    """Datos recolectados durante el onboarding del chatbot"""
    companyName: str
    rfc: str
    industry: str
    annualRevenue: str
    mainServices: str
    contactEmail: EmailStr
    documents: Optional[List[str]] = []
    
    @field_validator('rfc')
    @classmethod
    def validate_rfc(cls, v):
        pattern = r'^[A-ZÑ&]{3,4}[0-9]{6}[A-Z0-9]{3}$'
        if not re.match(pattern, v.upper().strip()):
            raise ValueError('RFC inválido. Formato esperado: 3-4 letras + 6 dígitos + 3 caracteres alfanuméricos')
        return v.upper().strip()
    
    @field_validator('companyName', 'industry', 'annualRevenue', 'mainServices')
    @classmethod
    def validate_not_empty(cls, v):
        if not v or not v.strip():
            raise ValueError('Este campo es obligatorio')
        return v.strip()


class OnboardingResponse(BaseModel):
    success: bool
    company_id: str
    company_name: str
    message: str
    error: Optional[str] = None


async def get_current_user_for_onboarding(
    credentials: HTTPAuthorizationCredentials = Depends(security)
) -> dict:
    """Obtiene el usuario actual del token JWT"""
    if not credentials:
        raise HTTPException(status_code=401, detail="No autenticado. Inicia sesión primero.")
    
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub") or payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Token inválido")
        
        user = await user_service.get_user_by_id(user_id)
        if not user:
            raise HTTPException(status_code=401, detail="Usuario no encontrado")
        
        return user.to_dict()
        
    except jose_exceptions.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expirado")
    except jose_exceptions.JWTError:
        raise HTTPException(status_code=401, detail="Token inválido")


async def create_company_in_db(data: OnboardingData, user_id: str) -> str:
    """Crea la empresa en PostgreSQL y retorna el ID"""
    from services.user_db import get_session, engine
    from sqlalchemy import text
    
    company_id = str(uuid.uuid4())
    
    async with engine.begin() as conn:
        await conn.execute(
            text("""
                INSERT INTO companies (id, name, rfc, industry, annual_revenue, main_services, contact_email, created_by, is_active, created_at)
                VALUES (:id, :name, :rfc, :industry, :annual_revenue, :main_services, :contact_email, :created_by, true, NOW())
            """),
            {
                "id": company_id,
                "name": data.companyName,
                "rfc": data.rfc,
                "industry": data.industry,
                "annual_revenue": data.annualRevenue,
                "main_services": data.mainServices,
                "contact_email": data.contactEmail,
                "created_by": user_id
            }
        )
    
    return company_id


async def add_company_to_user(user_id: str, company_id: str) -> bool:
    """Agrega la empresa a la lista de empresas permitidas del usuario"""
    user = await user_service.get_user_by_id(user_id)
    if not user:
        return False
    
    current_companies = []
    if user.allowed_companies:
        try:
            current_companies = json.loads(user.allowed_companies)
        except:
            current_companies = [c.strip() for c in user.allowed_companies.split(',') if c.strip()]
    
    if company_id not in current_companies:
        current_companies.append(company_id)
    
    return await user_service.update_user_allowed_companies(user_id, current_companies)


@onboarding_router.post("/complete", response_model=OnboardingResponse)
async def complete_onboarding(
    data: OnboardingData,
    current_user: dict = Depends(get_current_user_for_onboarding)
):
    """
    Completa el proceso de onboarding:
    1. Valida los datos recolectados
    2. Crea la empresa en PostgreSQL
    3. Asocia el usuario a la empresa
    4. Retorna el ID de la empresa creada
    """
    user_id = current_user.get("user_id")
    
    try:
        from services.user_db import get_session, engine
        from sqlalchemy import text
        
        async with engine.begin() as conn:
            result = await conn.execute(
                text("SELECT id FROM companies WHERE rfc = :rfc"),
                {"rfc": data.rfc}
            )
            existing = result.fetchone()
            if existing:
                raise HTTPException(
                    status_code=400,
                    detail=f"Ya existe una empresa registrada con el RFC {data.rfc}"
                )
        
        company_id = await create_company_in_db(data, user_id)
        
        association_success = await add_company_to_user(user_id, company_id)
        
        if not association_success:
            async with engine.begin() as conn:
                await conn.execute(
                    text("DELETE FROM companies WHERE id = :id"),
                    {"id": company_id}
                )
            raise HTTPException(
                status_code=500,
                detail="Error al asociar la empresa con tu usuario. Por favor, intenta de nuevo."
            )
        
        logger.info(f"Onboarding completado: Empresa {data.companyName} (ID: {company_id}) creada por usuario {user_id}")
        
        return OnboardingResponse(
            success=True,
            company_id=company_id,
            company_name=data.companyName,
            message=f"Empresa {data.companyName} registrada exitosamente"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en onboarding: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error al crear la empresa: {str(e)}"
        )


import asyncpg
from pathlib import Path
import hashlib

DATABASE_URL = os.environ.get('DATABASE_URL', '')
UPLOADS_BASE_PATH = Path("backend/uploads/clientes")


class ClienteIngestData(BaseModel):
    """Datos del cliente/proveedor para ingestar al expediente"""
    nombre: str
    rfc: Optional[str] = None
    razon_social: Optional[str] = None
    direccion: Optional[str] = None
    email: Optional[str] = None
    telefono: Optional[str] = None
    giro: Optional[str] = None
    sitio_web: Optional[str] = None
    regimen_fiscal: Optional[str] = None
    tipo_persona: Optional[str] = "persona_moral"
    actividad_economica: Optional[str] = None
    representante_legal: Optional[str] = None
    empresa_id: Optional[str] = None
    origen: Optional[str] = "archivo_chat"
    notas_internas: Optional[str] = None
    tags: Optional[List[str]] = []


class ClienteIngestResponse(BaseModel):
    success: bool
    cliente_id: Optional[int] = None
    cliente_uuid: Optional[str] = None
    nombre: Optional[str] = None
    message: str
    datos_guardados: Optional[dict] = None
    error: Optional[str] = None


class ExtractionRequest(BaseModel):
    """Request para extracción de datos con IA"""
    text: str
    document_type: Optional[str] = None
    existing_data: Optional[dict] = None


class ExtractionResponse(BaseModel):
    success: bool
    extracted_data: dict
    confidence: Optional[float] = None
    missing_fields: Optional[List[str]] = None
    message: str


async def get_db_connection():
    """Obtiene conexión a PostgreSQL"""
    if not DATABASE_URL:
        raise HTTPException(status_code=500, detail="DATABASE_URL no configurada")
    
    db_url = DATABASE_URL
    if db_url.startswith('postgres://'):
        db_url = db_url.replace('postgres://', 'postgresql://', 1)
    
    return await asyncpg.connect(db_url, timeout=30)


@onboarding_router.post("/ingest-client", response_model=ClienteIngestResponse)
async def ingest_client(
    data: ClienteIngestData,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """
    Ingesta información de un cliente/proveedor al expediente (tabla clientes).
    
    1. Valida los datos recibidos
    2. Verifica si ya existe un cliente con el mismo RFC
    3. Crea el registro en la tabla clientes de PostgreSQL
    4. Registra en historial de cambios
    5. Retorna el ID del cliente creado
    """
    user_context = await get_user_from_token(credentials)
    user_id = user_context.get('id') if user_context else None
    user_email = user_context.get('email') if user_context else 'sistema'
    
    try:
        if not data or not data.nombre:
            raise HTTPException(
                status_code=400,
                detail="El nombre del cliente es requerido"
            )

        conn = await get_db_connection()

        try:
            if data.rfc:
                rfc_upper = data.rfc.upper().strip()
                existing = await conn.fetchrow(
                    "SELECT id, nombre, rfc FROM clientes WHERE UPPER(rfc) = $1 AND activo = true",
                    rfc_upper
                )
                if existing:
                    return ClienteIngestResponse(
                        success=False,
                        cliente_id=existing['id'],
                        nombre=existing['nombre'],
                        message=f"Ya existe un cliente con RFC {rfc_upper}",
                        error="RFC_DUPLICADO"
                    )
            
            cliente_uuid = str(uuid.uuid4())
            now = datetime.utcnow()
            
            cliente_id = await conn.fetchval(
                """
                INSERT INTO clientes (
                    cliente_uuid, nombre, rfc, razon_social, direccion,
                    email, telefono, giro, sitio_web, regimen_fiscal,
                    tipo_persona, actividad_economica, empresa_id, origen,
                    notas_internas, tags, activo, estado,
                    created_at, updated_at, created_by
                )
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13, $14, $15, $16, $17, $18, $19, $20, $21)
                RETURNING id
                """,
                cliente_uuid,
                data.nombre,
                data.rfc.upper().strip() if data.rfc else None,
                data.razon_social or data.nombre,
                data.direccion,
                data.email,
                data.telefono,
                data.giro,
                data.sitio_web,
                data.regimen_fiscal,
                data.tipo_persona or "persona_moral",
                data.actividad_economica,
                data.empresa_id,
                data.origen or "archivo_chat",
                data.notas_internas,
                data.tags or [],
                True,
                "pendiente",
                now,
                now,
                user_id
            )
            
            await conn.execute(
                """
                INSERT INTO clientes_historial (
                    cliente_id, tipo_cambio, campo_modificado,
                    valor_nuevo, descripcion, origen, agente_id, created_at
                )
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8)
                """,
                cliente_id,
                "cliente_creado",
                "expediente",
                data.nombre,
                f"Cliente '{data.nombre}' creado via ARCHIVO agent",
                "archivo_routes",
                user_email,
                now
            )
            
            logger.info(f"Cliente {data.nombre} (ID: {cliente_id}) creado exitosamente")
            
            return ClienteIngestResponse(
                success=True,
                cliente_id=cliente_id,
                cliente_uuid=cliente_uuid,
                nombre=data.nombre,
                message=f"Cliente {data.nombre} creado exitosamente en el expediente",
                datos_guardados={
                    "nombre": data.nombre,
                    "rfc": data.rfc,
                    "razon_social": data.razon_social,
                    "email": data.email,
                    "telefono": data.telefono,
                    "giro": data.giro,
                    "direccion": data.direccion
                }
            )
            
        finally:
            await conn.close()
            
    except asyncpg.exceptions.UndefinedTableError as e:
        logger.error(f"Tabla clientes no existe: {e}")
        raise HTTPException(
            status_code=503,
            detail="La tabla clientes no existe en la base de datos. Contacte al administrador."
        )
    except asyncpg.exceptions.IntegrityConstraintViolationError as e:
        logger.error(f"Constraint violation: {e}")
        raise HTTPException(
            status_code=400,
            detail="Datos duplicados o constraint violation. Verifique los datos ingresados."
        )
    except Exception as e:
        logger.error(f"Error ingesting client: {e}")
        if "connection" in str(e).lower() or "timeout" in str(e).lower():
            status_code = 503
        else:
            status_code = 400
        raise HTTPException(
            status_code=status_code,
            detail=f"Error al crear el cliente: {str(e)}"
        )


@onboarding_router.post("/upload-document/{cliente_id}")
async def upload_client_document(
    cliente_id: int,
    file: UploadFile = File(...),
    tipo_documento: Optional[str] = Form(None),
    categoria: Optional[str] = Form(None),
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """
    Sube un documento y lo vincula a un cliente en clientes_documentos.
    
    1. Valida el archivo y el cliente
    2. Extrae texto del documento
    3. Guarda el archivo en disco
    4. Crea registro en clientes_documentos
    5. Registra en historial
    """
    user_context = await get_user_from_token(credentials)
    user_id = user_context.get('id') if user_context else None
    user_email = user_context.get('email') if user_context else 'sistema'
    
    try:
        filename = file.filename or "unknown"
        file_type, is_valid = detect_file_type(filename, file.content_type or '')

        if not is_valid:
            raise HTTPException(
                status_code=400,
                detail=f"Tipo de archivo no permitido: {file.content_type}. Extensión: {filename.split('.')[-1] if '.' in filename else 'none'}"
            )

        file_bytes = await file.read()

        if len(file_bytes) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"El archivo excede el límite de 50MB"
            )

        conn = await get_db_connection()

        try:
            if not cliente_id:
                raise HTTPException(status_code=400, detail="cliente_id es requerido")

            cliente = await conn.fetchrow(
                "SELECT id, nombre FROM clientes WHERE id = $1 AND activo = true",
                cliente_id
            )
            if not cliente:
                raise HTTPException(status_code=404, detail="Cliente no encontrado")

            extracted_text = ""
            extraction_method = "none"
            
            if file_type == 'pdf':
                extracted_text, extraction_method = extract_text_from_pdf(file_bytes)
            elif file_type == 'image':
                extracted_text, extraction_method = extract_text_from_image(file_bytes)
            elif file_type in ['docx', 'doc']:
                extracted_text, extraction_method = extract_text_from_docx(file_bytes)
            
            classification = archivo_service.classify_document_by_name(
                file_name=file.filename if file.filename else "documento",
                file_type=file.content_type if file.content_type else "application/octet-stream"
            )
            detected_type = tipo_documento or (classification.get("classification") if classification else "otro")
            
            hash_contenido = hashlib.sha256(file_bytes).hexdigest()
            
            existing_hash = await conn.fetchrow(
                """
                SELECT id, nombre_archivo FROM clientes_documentos 
                WHERE cliente_id = $1 AND hash_contenido = $2 AND activo = true
                """,
                cliente_id, hash_contenido
            )
            if existing_hash:
                return {
                    "success": False,
                    "message": "Este documento ya fue subido anteriormente",
                    "documento_existente": existing_hash['nombre_archivo']
                }
            
            upload_path = UPLOADS_BASE_PATH / str(cliente_id)
            upload_path.mkdir(parents=True, exist_ok=True)
            
            file_ext = Path(file.filename).suffix
            file_uuid = str(uuid.uuid4())
            saved_filename = f"{file_uuid}{file_ext}"
            file_path = upload_path / saved_filename
            
            file_path.write_bytes(file_bytes)
            
            now = datetime.utcnow()
            documento_uuid = str(uuid.uuid4())
            
            metadata_extraida = None
            if extracted_text:
                extracted_data = await archivo_service.extract_data_from_text(extracted_text)
                if extracted_data:
                    metadata_extraida = json.dumps(extracted_data)
            
            documento_id = await conn.fetchval(
                """
                INSERT INTO clientes_documentos (
                    cliente_id, documento_uuid, nombre_archivo, nombre_original,
                    tipo_mime, tamanio_bytes, ruta_archivo, hash_contenido,
                    version, es_version_actual, tipo_documento, categoria,
                    metadata_extraida, resumen_ia, procesado, origen,
                    activo, created_at, updated_at, created_by
                )
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13, $14, $15, $16, $17, $18, $19, $20)
                RETURNING id
                """,
                cliente_id,
                documento_uuid,
                saved_filename,
                file.filename,
                file.content_type,
                len(file_bytes),
                str(file_path),
                hash_contenido,
                1,
                True,
                detected_type,
                categoria,
                metadata_extraida,
                classification.get("summary"),
                True,
                "archivo_upload",
                True,
                now,
                now,
                user_id
            )
            
            await conn.execute(
                """
                INSERT INTO clientes_historial (
                    cliente_id, tipo_cambio, campo_modificado,
                    valor_nuevo, descripcion, origen, agente_id, created_at
                )
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8)
                """,
                cliente_id,
                "documento_subido",
                "documentos",
                file.filename,
                f"Documento '{file.filename}' subido al expediente ({detected_type})",
                "archivo_routes",
                user_email,
                now
            )
            
            logger.info(f"Documento {file.filename} subido para cliente {cliente_id}")
            
            return {
                "success": True,
                "documento_id": documento_id,
                "documento_uuid": documento_uuid,
                "nombre_archivo": file.filename,
                "tipo_documento": detected_type,
                "clasificacion": classification.get("classification_name"),
                "tamanio_bytes": len(file_bytes),
                "texto_extraido": len(extracted_text) > 0,
                "palabras_extraidas": len(extracted_text.split()) if extracted_text else 0,
                "metodo_extraccion": extraction_method,
                "message": f"Documento subido y vinculado al cliente {cliente['nombre']}"
            }
            
        finally:
            await conn.close()
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        if "not found" in str(e).lower():
            status_code = 404
        elif "connection" in str(e).lower() or "timeout" in str(e).lower():
            status_code = 503
        else:
            status_code = 400
        raise HTTPException(
            status_code=status_code,
            detail=f"Error al subir documento: {str(e)}"
        )


@onboarding_router.post("/extract-client-data", response_model=ExtractionResponse)
async def extract_client_data_ai(
    data: ExtractionRequest,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """
    Usa IA para extraer datos estructurados de cliente de texto o documentos.
    
    Extrae:
    - RFC (validado)
    - Razón Social / Nombre
    - Dirección fiscal
    - Régimen fiscal
    - Email, teléfono
    - Representante legal
    - Giro/actividad económica
    """
    try:
        extracted = await archivo_service.extract_data_from_text(data.text)
        
        if data.existing_data:
            for key, value in data.existing_data.items():
                if value and not extracted.get(key):
                    extracted[key] = value
        
        if extracted.get('rfc'):
            validation = archivo_service.validate_rfc(extracted['rfc'])
            extracted['rfc_valido'] = validation.get('valid', False)
            extracted['tipo_persona'] = validation.get('tipo', 'desconocido')
        
        campos_requeridos = ['nombre', 'razon_social', 'rfc', 'email', 'direccion']
        campos_faltantes = [c for c in campos_requeridos if not extracted.get(c)]
        
        datos_encontrados = sum(1 for v in extracted.values() if v)
        confidence = min(1.0, datos_encontrados / len(campos_requeridos))
        
        return ExtractionResponse(
            success=True,
            extracted_data=extracted,
            confidence=confidence,
            missing_fields=campos_faltantes if campos_faltantes else None,
            message=f"Se extrajeron {datos_encontrados} campos. Faltan: {', '.join(campos_faltantes) if campos_faltantes else 'ninguno'}"
        )
        
    except Exception as e:
        logger.error(f"AI extraction error: {e}")
        return ExtractionResponse(
            success=False,
            extracted_data={},
            confidence=0,
            missing_fields=None,
            message=f"Error en extracción: {str(e)}"
        )


@onboarding_router.post("/analyze-and-ingest")
async def analyze_and_ingest_document(
    file: UploadFile = File(...),
    auto_create: bool = Form(False),
    empresa_id: Optional[str] = Form(None),
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """
    Flujo completo: analiza documento, extrae datos y opcionalmente crea el cliente.
    
    1. Lee y valida el documento
    2. Extrae texto
    3. Usa IA para extraer datos del cliente
    4. Si auto_create=True y hay suficientes datos, crea el cliente
    5. Retorna datos extraídos y resultado de creación
    """
    user_context = await get_user_from_token(credentials)
    user_id = user_context.get('id') if user_context else None
    user_email = user_context.get('email') if user_context else 'sistema'
    
    try:
        filename = file.filename or "unknown"
        file_type, is_valid = detect_file_type(filename, file.content_type or '')

        if not is_valid:
            raise HTTPException(
                status_code=400,
                detail=f"Tipo de archivo no permitido: {file.content_type}. Extensión: {filename.split('.')[-1] if '.' in filename else 'none'}"
            )

        file_bytes = await file.read()

        if len(file_bytes) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail="El archivo excede el límite de 50MB"
            )

        extracted_text = ""
        extraction_method = "none"

        if file_type == 'pdf':
            extracted_text, extraction_method = extract_text_from_pdf(file_bytes)
        elif file_type == 'image':
            extracted_text, extraction_method = extract_text_from_image(file_bytes)
        elif file_type == 'text':
            extracted_text = file_bytes.decode('utf-8', errors='ignore')[:30000]
            extraction_method = "plain-text"
        elif file_type in ['docx', 'doc']:
            extracted_text, extraction_method = extract_text_from_docx(file_bytes)

        if not extracted_text:
            return {
                "success": False,
                "message": "No se pudo extraer texto del documento",
                "extraction_method": extraction_method
            }

        extracted_data = await archivo_service.extract_data_from_text(extracted_text)

        if extracted_data.get('rfc'):
            validation = archivo_service.validate_rfc(extracted_data['rfc'])
            extracted_data['rfc_valido'] = validation.get('valid', False)
            extracted_data['tipo_persona'] = validation.get('tipo')

        classification = archivo_service.classify_document_by_name(
            file_name=filename,
            file_type=file.content_type or ''
        )
        
        nombre = extracted_data.get('razon_social') or extracted_data.get('nombre')
        has_minimum = bool(nombre)
        
        cliente_creado = None
        
        if auto_create and has_minimum:
            try:
                conn = await get_db_connection()
                try:
                    if extracted_data.get('rfc'):
                        rfc_upper = extracted_data['rfc'].upper().strip()
                        existing = await conn.fetchrow(
                            "SELECT id, nombre FROM clientes WHERE UPPER(rfc) = $1 AND activo = true",
                            rfc_upper
                        )
                        if existing:
                            cliente_creado = {
                                "success": False,
                                "message": f"Ya existe cliente con RFC {rfc_upper}",
                                "cliente_id": existing['id'],
                                "nombre": existing['nombre']
                            }
                    
                    if not cliente_creado:
                        cliente_uuid = str(uuid.uuid4())
                        now = datetime.utcnow()
                        
                        cliente_id = await conn.fetchval(
                            """
                            INSERT INTO clientes (
                                cliente_uuid, nombre, rfc, razon_social, direccion,
                                email, telefono, giro, regimen_fiscal, tipo_persona,
                                empresa_id, origen, activo, estado,
                                created_at, updated_at, created_by
                            )
                            VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13, $14, $15, $16, $17)
                            RETURNING id
                            """,
                            cliente_uuid,
                            nombre,
                            extracted_data.get('rfc', '').upper().strip() if extracted_data.get('rfc') else None,
                            extracted_data.get('razon_social'),
                            extracted_data.get('direccion') or extracted_data.get('domicilio'),
                            extracted_data.get('email'),
                            extracted_data.get('telefono'),
                            extracted_data.get('giro') or extracted_data.get('actividad'),
                            extracted_data.get('regimen_fiscal'),
                            extracted_data.get('tipo_persona', 'persona_moral'),
                            empresa_id,
                            "archivo_auto_ingest",
                            True,
                            "pendiente",
                            now,
                            now,
                            user_id
                        )
                        
                        hash_contenido = hashlib.sha256(file_bytes).hexdigest()
                        upload_path = UPLOADS_BASE_PATH / str(cliente_id)
                        upload_path.mkdir(parents=True, exist_ok=True)
                        
                        file_ext = Path(file.filename).suffix
                        file_uuid = str(uuid.uuid4())
                        saved_filename = f"{file_uuid}{file_ext}"
                        file_path = upload_path / saved_filename
                        file_path.write_bytes(file_bytes)
                        
                        documento_uuid = str(uuid.uuid4())
                        
                        await conn.execute(
                            """
                            INSERT INTO clientes_documentos (
                                cliente_id, documento_uuid, nombre_archivo, nombre_original,
                                tipo_mime, tamanio_bytes, ruta_archivo, hash_contenido,
                                version, es_version_actual, tipo_documento,
                                metadata_extraida, procesado, origen,
                                activo, created_at, created_by
                            )
                            VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13, $14, $15, $16, $17)
                            """,
                            cliente_id,
                            documento_uuid,
                            saved_filename,
                            file.filename,
                            file.content_type,
                            len(file_bytes),
                            str(file_path),
                            hash_contenido,
                            1,
                            True,
                            classification.get("classification", "otro"),
                            json.dumps(extracted_data),
                            True,
                            "archivo_auto_ingest",
                            True,
                            now,
                            user_id
                        )
                        
                        await conn.execute(
                            """
                            INSERT INTO clientes_historial (
                                cliente_id, tipo_cambio, campo_modificado,
                                valor_nuevo, descripcion, origen, agente_id, created_at
                            )
                            VALUES ($1, $2, $3, $4, $5, $6, $7, $8)
                            """,
                            cliente_id,
                            "cliente_creado",
                            "expediente",
                            nombre,
                            f"Cliente creado automáticamente desde documento {file.filename}",
                            "archivo_auto_ingest",
                            user_email,
                            now
                        )
                        
                        cliente_creado = {
                            "success": True,
                            "cliente_id": cliente_id,
                            "cliente_uuid": cliente_uuid,
                            "nombre": nombre,
                            "message": f"Cliente {nombre} creado automáticamente"
                        }
                        
                        logger.info(f"Cliente {nombre} (ID: {cliente_id}) creado automáticamente desde documento")
                        
                finally:
                    await conn.close()
                    
            except Exception as e:
                logger.error(f"Error creating client from document: {e}")
                cliente_creado = {
                    "success": False,
                    "message": f"Error al crear cliente: {str(e)}"
                }
        
        campos_requeridos = ['razon_social', 'rfc', 'email', 'direccion']
        campos_faltantes = [c for c in campos_requeridos if not extracted_data.get(c)]
        
        return {
            "success": True,
            "document_info": {
                "filename": file.filename,
                "type": file.content_type,
                "size_bytes": len(file_bytes),
                "classification": classification.get("classification"),
                "classification_name": classification.get("classification_name")
            },
            "extraction": {
                "method": extraction_method,
                "text_length": len(extracted_text),
                "word_count": len(extracted_text.split())
            },
            "extracted_data": extracted_data,
            "missing_fields": campos_faltantes,
            "has_minimum_data": has_minimum,
            "cliente_creado": cliente_creado,
            "message": "Documento analizado exitosamente"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in analyze and ingest: {e}")
        if "file" in str(e).lower() or "not found" in str(e).lower():
            status_code = 404
        elif "timeout" in str(e).lower() or "connection" in str(e).lower():
            status_code = 503
        else:
            status_code = 400
        raise HTTPException(
            status_code=status_code,
            detail=f"Error procesando documento: {str(e)}"
        )


# ============================================================
# ARCHIVO ENTITY ROUTER - Para crear clientes/proveedores
# ============================================================
archivo_entity_router = APIRouter(prefix="/api/archivo", tags=["archivo-entity"])

# Import ProveedorService para MongoDB
try:
    from services.proveedor_service import ProveedorService
    proveedor_service = ProveedorService()
    PROVEEDOR_SERVICE_AVAILABLE = True
except ImportError:
    proveedor_service = None
    PROVEEDOR_SERVICE_AVAILABLE = False
    logger.warning("ProveedorService not available")


class CrearEntidadRequest(BaseModel):
    """Request para crear cliente o proveedor desde el chatbot de archivo"""
    tipo: str = "cliente"  # "cliente" o "proveedor"
    datos: dict
    email_contacto: Optional[str] = None
    archivos_ids: Optional[List[str]] = []
    empresa_id: Optional[str] = None  # Para proveedores, requerido


class CrearEntidadResponse(BaseModel):
    success: bool
    entity_id: Optional[str] = None
    entity_type: str
    nombre: Optional[str] = None
    message: str
    datos_guardados: Optional[dict] = None
    error: Optional[str] = None


@archivo_entity_router.post("/crear-entidad", response_model=CrearEntidadResponse)
async def crear_entidad(
    data: CrearEntidadRequest,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """
    Crea un cliente o proveedor desde el chatbot de ARCHIVO.

    Este endpoint es llamado por ChatbotArchivoRefactored cuando el usuario
    confirma los datos extraídos de los documentos.

    - Clientes: Se guardan en PostgreSQL (tabla clientes)
    - Proveedores: Se guardan en MongoDB via ProveedorService

    Args:
        data: Datos de la entidad a crear
        - tipo: "cliente" o "proveedor"
        - datos: Diccionario con campos extraídos (rfc, razon_social, nombre, etc.)
        - email_contacto: Email de contacto
        - archivos_ids: IDs de archivos ya subidos para vincular
        - empresa_id: ID de empresa (requerido para proveedores)

    Returns:
        CrearEntidadResponse con el resultado de la operación
    """
    user_context = await get_user_from_token(credentials)
    if not user_context:
        raise HTTPException(status_code=401, detail="No autenticado. Inicia sesión primero.")

    user_id = user_context.get('id') or str(uuid.uuid4())
    user_email = user_context.get('email', 'sistema')

    try:
        datos = data.datos or {}

        # Obtener nombre/razón social
        nombre = datos.get('razon_social') or datos.get('nombre') or datos.get('nombre_comercial')
        if not nombre:
            raise HTTPException(
                status_code=400,
                detail="Se requiere al menos el nombre o razón social de la entidad"
            )

        # Obtener RFC (opcional pero recomendado para proveedores)
        rfc = datos.get('rfc', '').upper().strip() if datos.get('rfc') else None

        # Email de contacto
        email = data.email_contacto or datos.get('email')

        # ================================================================
        # CREAR PROVEEDOR (MongoDB via ProveedorService)
        # ================================================================
        if data.tipo == 'proveedor':
            if not PROVEEDOR_SERVICE_AVAILABLE or not proveedor_service:
                raise HTTPException(
                    status_code=503,
                    detail="El servicio de proveedores no está disponible"
                )

            # Empresa ID es requerido para proveedores
            empresa_id = data.empresa_id or datos.get('empresa_id')
            if not empresa_id:
                # Intentar obtener de localStorage/header
                raise HTTPException(
                    status_code=400,
                    detail="Se requiere empresa_id para crear proveedores. Selecciona una empresa primero."
                )

            # Verificar RFC duplicado
            if rfc and proveedor_service.verificar_rfc_existente(rfc, empresa_id):
                return CrearEntidadResponse(
                    success=False,
                    entity_type="proveedor",
                    nombre=nombre,
                    message=f"Ya existe un proveedor con RFC {rfc} en esta empresa",
                    error="RFC_DUPLICADO"
                )

            # Preparar datos para ProveedorService
            proveedor_data = {
                "razon_social": datos.get('razon_social') or nombre,
                "nombre_comercial": datos.get('nombre_comercial') or datos.get('nombre'),
                "rfc": rfc or "",
                "regimen_fiscal": datos.get('regimen_fiscal') or "General de Ley",
                "tipo_persona": datos.get('tipo_persona', 'moral'),
                "tipo_proveedor": datos.get('tipo_proveedor', 'SERVICIOS'),
                "sitio_web": datos.get('sitio_web'),
                "domicilio_fiscal": {
                    "calle": datos.get('direccion') or datos.get('direccion_fiscal') or datos.get('domicilio') or "",
                    "colonia": datos.get('colonia', ''),
                    "ciudad": datos.get('ciudad', ''),
                    "estado": datos.get('estado', ''),
                    "cp": datos.get('codigo_postal') or datos.get('cp', ''),
                    "pais": "México"
                },
                "contacto_principal": {
                    "nombre": datos.get('contacto_nombre', ''),
                    "email": email or "",
                    "telefono": datos.get('telefono', '')
                }
            }

            result = proveedor_service.create_proveedor(proveedor_data, empresa_id, user_id)

            if result.get("service_unavailable"):
                raise HTTPException(
                    status_code=503,
                    detail="El servicio de proveedores no está disponible temporalmente"
                )

            if not result.get("success"):
                raise HTTPException(
                    status_code=400,
                    detail=result.get("error", "Error creando proveedor")
                )

            logger.info(f"Proveedor {nombre} creado exitosamente por {user_email}")

            return CrearEntidadResponse(
                success=True,
                entity_id=result.get("proveedor_id"),
                entity_type="proveedor",
                nombre=nombre,
                message=f"Proveedor {nombre} creado exitosamente",
                datos_guardados={
                    "id": result.get("proveedor_id"),
                    "nombre": nombre,
                    "rfc": rfc,
                    "email": email,
                    "empresa_id": empresa_id,
                    "tipo": "proveedor"
                }
            )

        # ================================================================
        # CREAR CLIENTE (PostgreSQL)
        # ================================================================
        conn = await get_db_connection()

        try:
            # Verificar RFC duplicado para clientes
            if rfc:
                existing = await conn.fetchrow(
                    "SELECT id, nombre FROM clientes WHERE UPPER(rfc) = $1 AND activo = true",
                    rfc
                )

                if existing:
                    return CrearEntidadResponse(
                        success=False,
                        entity_type="cliente",
                        nombre=existing.get('nombre'),
                        message=f"Ya existe un cliente con RFC {rfc}",
                        error="RFC_DUPLICADO"
                    )

            entity_uuid = str(uuid.uuid4())
            now = datetime.utcnow()

            # Crear cliente en PostgreSQL
            entity_id = await conn.fetchval(
                """
                INSERT INTO clientes (
                    cliente_uuid, nombre, rfc, razon_social, direccion,
                    email, telefono, giro, regimen_fiscal, tipo_persona,
                    origen, activo, estado, created_at, updated_at, created_by
                )
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13, $14, $15, $16)
                RETURNING id
                """,
                entity_uuid,
                nombre,
                rfc,
                datos.get('razon_social') or nombre,
                datos.get('direccion') or datos.get('domicilio'),
                email,
                datos.get('telefono'),
                datos.get('giro') or datos.get('actividad_economica'),
                datos.get('regimen_fiscal'),
                datos.get('tipo_persona', 'persona_moral'),
                "archivo_chatbot",
                True,
                "activo",
                now,
                now,
                user_id
            )

            # Registrar en historial
            try:
                await conn.execute(
                    """
                    INSERT INTO clientes_historial (
                        cliente_id, tipo_cambio, campo_modificado,
                        valor_nuevo, descripcion, origen, agente_id, created_at
                    )
                    VALUES ($1, $2, $3, $4, $5, $6, $7, $8)
                    """,
                    entity_id,
                    "cliente_creado",
                    "expediente",
                    nombre,
                    f"Cliente '{nombre}' creado via ARCHIVO chatbot",
                    "archivo_chatbot",
                    user_email,
                    now
                )
            except Exception as hist_err:
                logger.debug(f"Historial de cliente no registrado: {hist_err}")

            logger.info(f"Cliente {nombre} (ID: {entity_id}) creado exitosamente por {user_email}")

            return CrearEntidadResponse(
                success=True,
                entity_id=str(entity_id),
                entity_type="cliente",
                nombre=nombre,
                message=f"Cliente {nombre} creado exitosamente",
                datos_guardados={
                    "id": entity_id,
                    "uuid": entity_uuid,
                    "nombre": nombre,
                    "rfc": rfc,
                    "email": email,
                    "tipo": "cliente"
                }
            )

        finally:
            await conn.close()

    except HTTPException:
        raise
    except asyncpg.exceptions.UndefinedTableError as e:
        logger.error(f"Tabla no existe: {e}")
        raise HTTPException(
            status_code=503,
            detail=f"La tabla de clientes no existe. Contacte al administrador."
        )
    except asyncpg.exceptions.IntegrityConstraintViolationError as e:
        logger.error(f"Constraint violation: {e}")
        raise HTTPException(
            status_code=400,
            detail="Datos duplicados o inválidos. Verifique la información."
        )
    except Exception as e:
        logger.error(f"Error creando entidad: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error al crear {data.tipo}: {str(e)}"
        )
