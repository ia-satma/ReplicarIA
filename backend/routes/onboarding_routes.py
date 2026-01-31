"""
Rutas de API para el sistema de onboarding inteligente con IA
Permite crear clientes/proveedores analizando documentos y buscando datos en web
"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, Depends
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, EmailStr
from typing import List, Optional, Dict, Any
import json
import logging
import uuid
import os
import tempfile
from datetime import datetime, timezone

import asyncpg
from jose import jwt, exceptions as jose_exceptions
from services.document_analyzer import document_analyzer
from services.web_search_service import web_search_service
from services.empresa_service import empresa_service
from services.user_db import user_service
from services.cliente_service import cliente_service
from services.documento_versionado_service import documento_versionado_service
from models.empresa import EmpresaCreate, IndustriaEnum
from repositories.empresa_repository import empresa_repository
from services.deep_research_service import deep_research_service
from services.pcloud_service import pcloud_service

DATABASE_URL = os.environ.get('DATABASE_URL', '')

try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    fitz = None
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/archivo", tags=["onboarding"])

security = HTTPBearer(auto_error=False)
SECRET_KEY: str = os.getenv("SECRET_KEY") or os.getenv("SESSION_SECRET") or os.getenv("JWT_SECRET_KEY") or "fallback-dev-key"
if SECRET_KEY == "fallback-dev-key":
    logger.error("CRITICAL: SESSION_SECRET or JWT_SECRET_KEY must be configured")
ALGORITHM = "HS256"


class WebSearchRequest(BaseModel):
    datos_actuales: Dict[str, Any]
    campos_faltantes: List[str]


class CrearEntidadRequest(BaseModel):
    tipo: str
    datos: Dict[str, Any]
    email_contacto: Optional[str] = ""
    archivos_ids: Optional[List[str]] = []


class InvestigarEmpresaRequest(BaseModel):
    """Request para investigar empresa automáticamente"""
    sitio_web: Optional[str] = None
    rfc: Optional[str] = None
    nombre: Optional[str] = None


async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Obtener usuario actual del token JWT o sesión OTP"""
    if not credentials:
        return None
    
    token = credentials.credentials
    
    # Try JWT first
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        return payload
    except (jwt.JWTError, jwt.ExpiredSignatureError, jwt.JWTClaimsError, Exception) as e:
        # JWT validation failed, will try OTP session next
        pass
    
    # Fallback to OTP session token
    try:
        from services.otp_auth_service import otp_auth_service
        session = await otp_auth_service.validate_session(token)
        if session and session.get("user"):
            user = session["user"]
            return {
                "user_id": user.get("id"),
                "email": user.get("email"),
                "full_name": user.get("nombre"),
                "empresa_id": user.get("empresa"),
                "company_id": user.get("empresa"),
                "role": user.get("rol", "user")
            }
    except Exception as e:
        logger.warning(f"OTP session validation failed: {e}")
    
    return None


@router.get("/status")
async def get_archivo_status():
    """Obtiene el estado del servicio ARCHIVO (onboarding chatbot)"""
    # Check if AI provider is actually configured
    ia_disponible = False
    ia_provider = None
    anthropic_status = "not_checked"
    openai_status = "not_checked"

    try:
        from services.document_analyzer import AI_PROVIDER
        ia_disponible = AI_PROVIDER is not None
        ia_provider = AI_PROVIDER
    except ImportError:
        pass

    # Check individual providers
    try:
        from services.anthropic_provider import is_configured as anthropic_configured
        anthropic_status = "configured" if anthropic_configured() else "not_configured"
    except ImportError:
        anthropic_status = "not_installed"

    try:
        from services.openai_provider import is_configured as openai_configured
        openai_status = "configured" if openai_configured() else "not_configured"
    except ImportError:
        openai_status = "not_installed"

    # Check env vars (don't expose actual keys)
    anthropic_key_set = bool(os.environ.get('ANTHROPIC_API_KEY', ''))
    openai_key_set = bool(os.environ.get('OPENAI_API_KEY', ''))

    return {
        "active": True,
        "status": "ready" if ia_disponible else "limited",
        "servicio": "ARCHIVO - Onboarding Inteligente",
        "version": "1.0.2",
        "capacidades": [
            "Análisis de documentos (PDF, DOCX, imágenes)",
            "Extracción automática de datos (RFC, razón social, dirección)",
            "Búsqueda web para completar información",
            "Creación de clientes y proveedores",
            "Validación de datos fiscales"
        ],
        "formatos_soportados": {
            "pdf": PYMUPDF_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "imagenes": True
        },
        "ia_disponible": ia_disponible,
        "ia_provider": ia_provider,
        "providers": {
            "anthropic": anthropic_status,
            "openai": openai_status,
            "anthropic_key_set": anthropic_key_set,
            "openai_key_set": openai_key_set
        },
        "warning": None if ia_disponible else "IA no configurada: Configure ANTHROPIC_API_KEY o OPENAI_API_KEY en las variables de entorno"
    }


@router.get("/diagnostico")
async def diagnostico_ia():
    """
    Endpoint de diagnóstico detallado para debugging de problemas con IA.
    Solo para desarrollo/debugging.
    """
    result = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "environment": os.environ.get('ENVIRONMENT', 'unknown'),
        "providers": {},
        "extraction": {},
        "env_vars": {}
    }

    # Check Anthropic
    try:
        from services.anthropic_provider import is_configured as anthropic_configured, ANTHROPIC_AVAILABLE
        result["providers"]["anthropic"] = {
            "package_installed": ANTHROPIC_AVAILABLE,
            "is_configured": anthropic_configured(),
            "api_key_set": bool(os.environ.get('ANTHROPIC_API_KEY', ''))
        }
    except Exception as e:
        result["providers"]["anthropic"] = {"error": str(e)}

    # Check OpenAI
    try:
        from services.openai_provider import is_configured as openai_configured, OPENAI_AVAILABLE
        result["providers"]["openai"] = {
            "package_installed": OPENAI_AVAILABLE,
            "is_configured": openai_configured(),
            "api_key_set": bool(os.environ.get('OPENAI_API_KEY', ''))
        }
    except Exception as e:
        result["providers"]["openai"] = {"error": str(e)}

    # Check document analyzer
    try:
        from services.document_analyzer import AI_PROVIDER, chat_fn
        result["providers"]["document_analyzer"] = {
            "ai_provider": AI_PROVIDER,
            "chat_fn_available": chat_fn is not None
        }
    except Exception as e:
        result["providers"]["document_analyzer"] = {"error": str(e)}

    # Extraction capabilities
    result["extraction"] = {
        "pdf": PYMUPDF_AVAILABLE,
        "docx": DOCX_AVAILABLE
    }

    # Relevant env vars (existence only, not values)
    env_keys_to_check = [
        'OPENAI_API_KEY', 'ANTHROPIC_API_KEY', 'ENVIRONMENT',
        'DATABASE_URL', 'MONGO_URL'
    ]
    for key in env_keys_to_check:
        val = os.environ.get(key, '')
        result["env_vars"][key] = "set" if val else "not_set"

    return result


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extrae texto de un PDF"""
    if not PYMUPDF_AVAILABLE or fitz is None:
        return "[PDF no soportado - instalar PyMuPDF]"
    
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text = text + str(page_text)
        doc.close()
        return text[:30000]
    except Exception as e:
        logger.error(f"Error extracting PDF: {e}")
        return ""


def extract_text_from_docx(file_content: bytes) -> str:
    """Extrae texto de un archivo DOCX incluyendo tablas"""
    logger.info(f"DOCX extraction starting, content size: {len(file_content)} bytes")

    # Method 1: Try with python-docx via BytesIO
    if DOCX_AVAILABLE and DocxDocument is not None:
        try:
            import io
            doc = DocxDocument(io.BytesIO(file_content))
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Extract tables (important for company data)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            text = "\n".join(text_parts)
            logger.info(f"DOCX BytesIO extracted: {len(text)} chars from {len(text_parts)} elements")
            if text.strip():
                return text[:30000]
        except Exception as e:
            logger.warning(f"DOCX BytesIO extraction failed: {e}")

    # Method 2: Try with temp file
    if DOCX_AVAILABLE and DocxDocument is not None:
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            doc = DocxDocument(tmp_path)
            text_parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            text = "\n".join(text_parts)
            os.unlink(tmp_path)
            logger.info(f"DOCX temp file extracted: {len(text)} chars from {len(text_parts)} elements")
            if text.strip():
                return text[:30000]
        except Exception as e2:
            logger.warning(f"DOCX temp file extraction failed: {e2}")
            try:
                os.unlink(tmp_path)
            except OSError:
                pass  # File cleanup failed, not critical

    # Method 3: Raw XML extraction from ZIP (DOCX is a ZIP with XML files)
    try:
        import zipfile
        import io
        import re

        with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
            # Read the main document.xml
            if 'word/document.xml' in zf.namelist():
                xml_content = zf.read('word/document.xml').decode('utf-8', errors='ignore')
                # Extract text between <w:t> tags
                text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                text = ' '.join(text_matches)
                # Clean up excessive whitespace
                text = re.sub(r'\s+', ' ', text).strip()
                logger.info(f"DOCX raw XML extracted: {len(text)} chars")
                if text:
                    return text[:30000]
    except Exception as e3:
        logger.warning(f"DOCX raw XML extraction failed: {e3}")

    # Method 4: Try to decode as plain text (in case it's not a real DOCX)
    try:
        text = file_content.decode('utf-8', errors='ignore')[:30000]
        # Check if it looks like text
        if text and len([c for c in text[:1000] if c.isalpha()]) > 100:
            logger.info(f"DOCX decoded as plain text: {len(text)} chars")
            return text
    except (UnicodeDecodeError, AttributeError):
        pass  # Not decodable as text

    logger.error("All DOCX extraction methods failed")
    return ""


async def crear_cliente_postgresql(
    datos: Dict[str, Any],
    empresa_id: Optional[str] = None,
    usuario_id: Optional[str] = None,
    origen: str = "chat"
) -> Optional[Dict[str, Any]]:
    """
    Crea un cliente en la tabla PostgreSQL 'clientes' con estado='pendiente'.
    Retorna el cliente creado con su ID.
    """
    if not DATABASE_URL:
        logger.warning("DATABASE_URL no configurada, no se puede crear cliente en PostgreSQL")
        return None
    
    db_url = DATABASE_URL
    if db_url.startswith('postgres://'):
        db_url = db_url.replace('postgres://', 'postgresql://', 1)
    
    try:
        conn = await asyncpg.connect(db_url)
        
        cliente_uuid = str(uuid.uuid4())
        now = datetime.utcnow()
        
        cliente_id = await conn.fetchval(
            """
            INSERT INTO clientes (
                cliente_uuid, nombre, rfc, razon_social, direccion, email, telefono,
                giro, sitio_web, regimen_fiscal, tipo_persona, actividad_economica,
                estado, origen, empresa_id, usuario_responsable_id, activo,
                created_at, updated_at, created_by
            )
            VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13, $14, $15, $16, $17, $18, $19, $20)
            RETURNING id
            """,
            cliente_uuid,
            datos.get("nombre") or datos.get("razon_social") or "Sin nombre",
            datos.get("rfc", "").upper() if datos.get("rfc") else None,
            datos.get("razon_social") or datos.get("nombre"),
            datos.get("direccion"),
            datos.get("email"),
            datos.get("telefono"),
            datos.get("giro"),
            datos.get("sitio_web"),
            datos.get("regimen_fiscal"),
            datos.get("tipo_persona"),
            datos.get("actividad_economica") or datos.get("giro"),
            "pendiente",
            origen,
            empresa_id,
            usuario_id,
            True,
            now,
            now,
            usuario_id
        )
        
        await conn.close()
        
        logger.info(f"Cliente creado en PostgreSQL: id={cliente_id}, uuid={cliente_uuid}")
        
        return {
            "id": cliente_id,
            "cliente_uuid": cliente_uuid,
            "nombre": datos.get("nombre") or datos.get("razon_social"),
            "rfc": datos.get("rfc"),
            "razon_social": datos.get("razon_social") or datos.get("nombre"),
            "estado": "pendiente",
            "origen": origen
        }
        
    except Exception as e:
        logger.error(f"Error creando cliente en PostgreSQL: {e}")
        return None


async def guardar_cliente_pcloud(
    datos: Dict[str, Any],
    cliente_id: str,
    tipo_entidad: str = "cliente"
) -> Optional[Dict[str, Any]]:
    """
    Guarda los datos del cliente en pCloud para referencia futura.
    Crea estructura: DEFENSE_FILES/{RFC}/{cliente_info.json}
    """
    try:
        # Check if pCloud is available
        if not pcloud_service.is_available():
            logger.info("pCloud not configured, skipping upload")
            return {"pcloud_path": "demo_mode", "pcloud_file_id": None}

        rfc = datos.get("rfc", "SIN_RFC").upper().replace(" ", "_")
        fecha = datetime.now().strftime("%Y-%m-%d")

        # Create client info document
        cliente_doc = {
            "id": cliente_id,
            "tipo": tipo_entidad,
            "fecha_creacion": fecha,
            "ultima_actualizacion": datetime.now().isoformat(),
            "datos": {
                "nombre": datos.get("nombre"),
                "razon_social": datos.get("razon_social"),
                "rfc": datos.get("rfc"),
                "direccion": datos.get("direccion"),
                "telefono": datos.get("telefono"),
                "email": datos.get("email"),
                "giro": datos.get("giro"),
                "regimen_fiscal": datos.get("regimen_fiscal"),
                "representante_legal": datos.get("representante_legal"),
                "sitio_web": datos.get("sitio_web"),
                "capital_social": datos.get("capital_social")
            },
            "fuentes": datos.get("fuentes", []),
            "proyectos": [],
            "versiones": [
                {
                    "version": "1.0",
                    "fecha": fecha,
                    "cambios": "Creación inicial"
                }
            ]
        }

        # Login to pCloud first
        login_result = pcloud_service.login()
        if not login_result.get("success"):
            logger.warning(f"pCloud login failed: {login_result.get('error')}")
            return None

        # Get DEFENSE_FILES folder ID
        from services.pcloud_service import AGENT_FOLDER_IDS
        defense_files_folder_id = AGENT_FOLDER_IDS.get("DEFENSE_FILES", 0)

        if not defense_files_folder_id:
            logger.warning("DEFENSE_FILES folder ID not found")
            return None

        # Create RFC subfolder if needed
        rfc_folder_result = pcloud_service.create_folder(defense_files_folder_id, rfc)
        rfc_folder_id = rfc_folder_result.get("folder_id", defense_files_folder_id)

        # Save to pCloud using the correct method signature: (folder_id, filename, content)
        content = json.dumps(cliente_doc, indent=2, ensure_ascii=False, default=str)
        filename = f"cliente_info_{cliente_id}.json"

        result = pcloud_service.upload_file(
            folder_id=rfc_folder_id,
            filename=filename,
            content=content.encode('utf-8')
        )

        if result and result.get("success"):
            folder_path = f"DEFENSE_FILES/{rfc}"
            logger.info(f"Client data saved to pCloud: {folder_path}/{filename}")
            return {
                "pcloud_path": f"{folder_path}/{filename}",
                "pcloud_file_id": result.get("file_id")
            }
        else:
            logger.warning(f"pCloud upload failed: {result}")
            return None

    except Exception as e:
        logger.error(f"Error saving client to pCloud: {e}")
        return None


class DocumentUploadRequest(BaseModel):
    cliente_id: int
    tipo_documento: Optional[str] = None
    categoria: Optional[str] = None
    subcategoria: Optional[str] = None


@router.post("/analizar")
async def analizar_documentos(
    files: List[UploadFile] = File(...),
    tipo_entidad: str = Form("cliente"),
    email_contacto: str = Form(""),
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """
    Analiza documentos subidos y extrae datos del cliente/proveedor usando IA
    """
    if not files:
        raise HTTPException(status_code=400, detail="No se proporcionaron archivos")

    documentos_procesados = []
    archivos_fallidos = []

    logger.info(f"Analizando {len(files)} archivos para {tipo_entidad}")

    for file in files:
        try:
            content = await file.read()
            texto = ""
            filename = file.filename or "unknown"
            content_type = file.content_type or ""

            logger.info(f"Procesando archivo: {filename} ({content_type}, {len(content)} bytes)")

            # Determine file type and extract text
            is_pdf = content_type == "application/pdf" or filename.lower().endswith('.pdf')
            is_docx = content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.lower().endswith('.docx')
            is_text = content_type.startswith("text/") if content_type else False

            if is_pdf:
                texto = extract_text_from_pdf(content)
                logger.info(f"PDF extraction result: {len(texto)} chars")
            elif is_docx:
                texto = extract_text_from_docx(content)
                logger.info(f"DOCX extraction result: {len(texto)} chars")
            elif is_text or filename.lower().endswith(('.txt', '.csv')):
                texto = content.decode('utf-8', errors='ignore')[:30000]
                logger.info(f"Text file decoded: {len(texto)} chars")
            else:
                # Try to decode as text anyway
                try:
                    texto = content.decode('utf-8', errors='ignore')[:30000]
                    logger.info(f"Generic decode: {len(texto)} chars")
                except (UnicodeDecodeError, AttributeError) as e:
                    logger.warning(f"Could not decode {filename}: {e}")
                    texto = ""

            if texto and texto.strip():
                documentos_procesados.append({
                    "nombre": filename,
                    "texto": texto,
                    "tipo": content_type
                })
                logger.info(f"Successfully processed: {filename}")
            else:
                archivos_fallidos.append(filename)
                logger.warning(f"No text extracted from: {filename}")

        except Exception as e:
            logger.error(f"Error processing file {file.filename}: {e}", exc_info=True)
            archivos_fallidos.append(file.filename or "unknown")
            continue

    if not documentos_procesados:
        error_msg = "No se pudo extraer texto de ningún documento"
        if archivos_fallidos:
            error_msg += f". Archivos fallidos: {', '.join(archivos_fallidos)}"
        logger.error(error_msg)
        raise HTTPException(status_code=400, detail=error_msg)

    try:
        datos = await document_analyzer.analizar_documentos(
            documentos=documentos_procesados,
            tipo_entidad=tipo_entidad
        )
    except Exception as e:
        logger.error(f"Document analysis exception: {e}")
        datos = {"error": str(e)}

    # Check if AI analysis failed (no provider configured)
    if not datos or datos.get("error"):
        logger.error(f"Document analysis failed: {datos.get('error') if datos else 'No response'}")
        # Return partial success - text was extracted but AI analysis failed
        return {
            "success": False,
            "error": datos.get("error") if datos else "Analysis service unavailable",
            "datos": datos or {},
            "archivos_procesados": len(documentos_procesados),
            "archivos_nombres": [d["nombre"] for d in documentos_procesados],
            "texto_extraido": True,
            "ia_disponible": False,
            "mensaje": "Se extrajo el texto de los documentos pero el análisis con IA falló. Verifique la configuración de ANTHROPIC_API_KEY o OPENAI_API_KEY."
        }

    if email_contacto and not datos.get("email"):
        # Null check
        if email_contacto and email_contacto.strip():
            datos["email"] = email_contacto
            if "fuentes" not in datos:
                datos["fuentes"] = []
            datos["fuentes"].append({
                "campo": "email",
                "fuente": "usuario",
                "confianza": 1.0
            })

    # Auto-trigger deep research if critical data is missing
    campos_criticos_faltantes = []
    if not datos.get("rfc"):
        campos_criticos_faltantes.append("rfc")
    if not datos.get("nombre") and not datos.get("razon_social"):
        campos_criticos_faltantes.append("nombre")
    if not datos.get("direccion"):
        campos_criticos_faltantes.append("direccion")

    datos_investigacion = {}
    if campos_criticos_faltantes and (datos.get("nombre") or datos.get("razon_social") or datos.get("rfc")):
        try:
            logger.info(f"Auto-triggering deep research for missing fields: {campos_criticos_faltantes}")
            investigacion = await deep_research_service.investigar_empresa(
                nombre=datos.get("nombre") or datos.get("razon_social"),
                rfc=datos.get("rfc"),
                documento=documentos_procesados[0]["texto"] if documentos_procesados else None
            )

            if investigacion.get("data"):
                datos_investigacion = investigacion["data"]
                # Merge found data
                for campo in campos_criticos_faltantes:
                    if datos_investigacion.get(campo) and not datos.get(campo):
                        datos[campo] = datos_investigacion[campo]
                        if "fuentes" not in datos:
                            datos["fuentes"] = []
                        datos["fuentes"].append({
                            "campo": campo,
                            "fuente": "deep_research",
                            "confianza": investigacion.get("field_confidence", {}).get(campo, {}).get("confidence", 70) / 100
                        })
                        logger.info(f"Deep research found: {campo} = {datos[campo]}")

                # Update missing fields list
                datos["datosFaltantes"] = [f for f in datos.get("datosFaltantes", []) if f not in datos_investigacion]

        except Exception as e:
            logger.warning(f"Deep research failed (non-critical): {e}")

    return {
        "success": True,
        "datos": datos,
        "archivos_procesados": len(documentos_procesados),
        "archivos_nombres": [d["nombre"] for d in documentos_procesados],
        "investigacion_automatica": bool(datos_investigacion),
        "campos_investigados": list(datos_investigacion.keys()) if datos_investigacion else []
    }


@router.post("/buscar-web")
async def buscar_datos_web(request: WebSearchRequest, credentials: HTTPAuthorizationCredentials = Depends(security)):
    """
    Busca datos faltantes de una empresa en fuentes públicas de internet
    """
    nombre = request.datos_actuales.get("nombre") or request.datos_actuales.get("razon_social")
    rfc = request.datos_actuales.get("rfc")
    
    if not nombre and not rfc:
        return {
            "datos_encontrados": {},
            "fuentes_nuevas": [],
            "aun_faltantes": request.campos_faltantes
        }
    
    resultado = await web_search_service.buscar_datos_empresa(
        nombre=nombre,
        rfc=rfc,
        campos_faltantes=request.campos_faltantes
    )
    
    return resultado


@router.post("/investigar")
async def investigar_empresa(
    request: InvestigarEmpresaRequest,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """
    Deep Research: Investiga automáticamente una empresa a partir de input mínimo.
    
    Acepta:
    - sitio_web: URL del sitio web de la empresa
    - rfc: RFC de la empresa
    - nombre: Nombre o razón social
    
    Retorna datos auto-completados con niveles de confianza por campo.
    """
    if not request.sitio_web and not request.rfc and not request.nombre:
        raise HTTPException(
            status_code=400,
            detail="Debe proporcionar al menos un dato: sitio_web, rfc o nombre"
        )
    
    logger.info(f"Iniciando Deep Research: web={request.sitio_web}, rfc={request.rfc}, nombre={request.nombre}")

    try:
        resultado = await deep_research_service.investigar_empresa(
            sitio_web=request.sitio_web,
            rfc=request.rfc,
            nombre=request.nombre
        )

        if not resultado:
            raise HTTPException(status_code=503, detail="Research service returned no results")

        return resultado

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en Deep Research: {e}")
        raise HTTPException(
            status_code=503,
            detail="Investigation service temporarily unavailable. Please try again later."
        )


@router.post("/crear-entidad")
async def crear_entidad(request: CrearEntidadRequest, credentials: HTTPAuthorizationCredentials = Depends(security)):
    """
    Crea un nuevo cliente o proveedor con los datos recopilados.
    - Si el usuario está autenticado: crea el cliente dentro de su empresa (tenant)
    - Si no hay auth o es admin creando tenant: crea una empresa nueva
    """
    datos = request.datos
    
    if not datos.get("rfc") and not datos.get("nombre") and not datos.get("razon_social"):
        raise HTTPException(
            status_code=400, 
            detail="Se requiere al menos el RFC o nombre de la empresa"
        )
    
    current_user = await get_current_user(credentials)
    user_empresa_id = None
    user_id = None
    is_admin = False
    
    if current_user:
        user_empresa_id = current_user.get("empresa_id") or current_user.get("company_id")
        user_id = current_user.get("user_id") or current_user.get("sub")
        role = current_user.get("role", "").lower()
        is_admin = role in ["admin", "superadmin", "platform_admin"]
    
    try:
        if current_user and not is_admin and not user_empresa_id:
            raise HTTPException(
                status_code=403,
                detail="Usuario sin empresa asignada. Contacte al administrador."
            )
        
        if current_user and not is_admin and request.tipo != "cliente":
            raise HTTPException(
                status_code=403,
                detail="Solo administradores pueden crear nuevos tenants. Use tipo='cliente'."
            )
        
        if not is_admin and not current_user:
            raise HTTPException(
                status_code=401,
                detail="Autenticación requerida para crear entidades."
            )
        
        if request.tipo == "cliente" and user_empresa_id:
            cliente_data = {
                "razon_social": datos.get("razon_social") or datos.get("nombre") or "Sin nombre",
                "nombre_comercial": datos.get("nombre") or datos.get("razon_social"),
                "rfc": datos.get("rfc", "").upper() if datos.get("rfc") else None,
                "direccion_fiscal": datos.get("direccion"),
                "ciudad": datos.get("ciudad"),
                "estado": datos.get("estado"),
                "codigo_postal": datos.get("codigo_postal"),
                "regimen_fiscal": datos.get("regimen_fiscal"),
                "actividad_economica": datos.get("giro"),
                "giro": datos.get("giro"),
                "representante_legal": datos.get("representante_legal"),
                "sitio_web": datos.get("sitio_web"),
                "telefono_principal": datos.get("telefono"),
                "email_principal": request.email_contacto or datos.get("email"),
                "status": "activo"
            }
            
            cliente = await cliente_service.create_cliente(
                cliente_data={k: v for k, v in cliente_data.items() if v is not None},
                empresa_id=user_empresa_id,
                creado_por=user_id
            )
            
            pg_datos = {
                "nombre": datos.get("nombre") or datos.get("razon_social"),
                "razon_social": datos.get("razon_social") or datos.get("nombre"),
                "rfc": datos.get("rfc"),
                "direccion": datos.get("direccion"),
                "email": request.email_contacto or datos.get("email"),
                "telefono": datos.get("telefono"),
                "giro": datos.get("giro"),
                "sitio_web": datos.get("sitio_web"),
                "regimen_fiscal": datos.get("regimen_fiscal"),
            }
            pg_cliente = await crear_cliente_postgresql(
                datos=pg_datos,
                empresa_id=user_empresa_id,
                usuario_id=user_id,
                origen="chat"
            )
            
            pg_cliente_id = pg_cliente.get("id") if pg_cliente else None
            
            logger.info(f"Created cliente {cliente['id']} for empresa {user_empresa_id}, PostgreSQL ID: {pg_cliente_id}")

            # Save to pCloud for persistent storage
            pcloud_result = await guardar_cliente_pcloud(
                datos=datos,
                cliente_id=str(cliente["id"]),
                tipo_entidad="cliente"
            )

            return {
                "success": True,
                "cliente_id": cliente["id"],
                "pg_cliente_id": pg_cliente_id,
                "empresa_id": user_empresa_id,
                "tipo": "cliente",
                "company_name": cliente.get("razon_social"),
                "mensaje": "Cliente creado exitosamente",
                "datos_guardados": {
                    "nombre": cliente.get("razon_social"),
                    "rfc": cliente.get("rfc"),
                    "giro": cliente.get("giro")
                },
                "pcloud": pcloud_result
            }
        
        industria = IndustriaEnum.OTRO
        giro = datos.get("giro", "").lower() if datos.get("giro") else ""
        
        if "tecnolog" in giro or "software" in giro or "sistemas" in giro:
            industria = IndustriaEnum.TECNOLOGIA
        elif "construc" in giro or "inmobili" in giro:
            industria = IndustriaEnum.CONSTRUCCION
        elif "manufactur" in giro or "fabric" in giro:
            industria = IndustriaEnum.MANUFACTURA
        elif "comerci" in giro or "venta" in giro:
            industria = IndustriaEnum.COMERCIO
        elif "servicio" in giro or "consultor" in giro:
            industria = IndustriaEnum.SERVICIOS_PROFESIONALES
        elif "salud" in giro or "hospital" in giro or "medic" in giro:
            industria = IndustriaEnum.SALUD
        elif "transporte" in giro or "logistic" in giro:
            industria = IndustriaEnum.TRANSPORTE_LOGISTICA
        
        rfc_value = datos.get("rfc", "").upper() if datos.get("rfc") else "PENDIENTE000"
        empresa_data = EmpresaCreate(
            nombre_comercial=datos.get("nombre") or datos.get("razon_social") or "Sin nombre",
            razon_social=datos.get("razon_social") or datos.get("nombre") or "Sin razón social",
            rfc=rfc_value,
            industria=industria,
            sub_industria=datos.get("giro")
        )
        
        empresa = await empresa_service.crear_empresa(empresa_data)
        empresa_id = empresa.id
        
        campos_extra = {}
        if datos.get("direccion"):
            campos_extra["direccion"] = datos["direccion"]
        if datos.get("telefono"):
            campos_extra["telefono"] = datos["telefono"]
        if datos.get("sitio_web"):
            campos_extra["sitio_web"] = datos["sitio_web"]
        if datos.get("representante_legal"):
            campos_extra["representante_legal"] = datos["representante_legal"]
        if datos.get("capital_social"):
            campos_extra["capital_social"] = datos["capital_social"]
        if datos.get("regimen_fiscal"):
            campos_extra["regimen_fiscal"] = datos["regimen_fiscal"]
        
        if campos_extra:
            await empresa_repository.update(empresa_id, campos_extra)
        
        usuario_id = None
        if request.email_contacto:
            try:
                existing_user = await user_service.get_user_by_email(request.email_contacto)
                if existing_user:
                    user_dict = existing_user.to_dict() if hasattr(existing_user, 'to_dict') else {}
                    usuario_id = user_dict.get("user_id") if user_dict else getattr(existing_user, 'id', None)
                else:
                    nuevo_usuario = await user_service.create_user({
                        "user_id": str(uuid.uuid4()),
                        "email": request.email_contacto,
                        "full_name": datos.get("representante_legal") or datos.get("nombre") or "Usuario",
                        "password_hash": None,
                        "role": "cliente" if request.tipo == "cliente" else "proveedor",
                        "company": empresa_id,
                        "is_active": False,
                        "approval_status": "pending"
                    })
                    usuario_id = nuevo_usuario.id if nuevo_usuario else None
            except Exception as e:
                logger.warning(f"Could not create user: {e}")
        
        logger.info(f"Created {request.tipo}: {empresa_id} with email {request.email_contacto}")
        
        return {
            "success": True,
            "company_id": empresa_id,
            "empresa_id": empresa_id,
            "usuario_id": usuario_id,
            "tipo": request.tipo,
            "company_name": empresa.nombre_comercial,
            "mensaje": f"{request.tipo.capitalize()} creado exitosamente",
            "datos_guardados": {
                "nombre": empresa.nombre_comercial,
                "rfc": empresa.rfc,
                "industria": empresa.industria.value if hasattr(empresa.industria, 'value') else str(empresa.industria)
            }
        }
        
    except HTTPException:
        raise
    except ValueError as e:
        logger.error(f"Validation error creating entity: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error creating entity: {e}")
        raise HTTPException(status_code=503, detail="Unable to create entity. Service temporarily unavailable.")


@router.get("/tipos-entidad")
async def get_tipos_entidad():
    """Retorna los tipos de entidad disponibles"""
    return {
        "tipos": [
            {"id": "cliente", "nombre": "Cliente", "descripcion": "Empresa que contrata servicios"},
            {"id": "proveedor", "nombre": "Proveedor", "descripcion": "Empresa que proporciona servicios"}
        ]
    }


@router.post("/documentos")
async def subir_documentos_cliente(
    files: List[UploadFile] = File(...),
    cliente_id: int = Form(...),
    tipo_documento: Optional[str] = Form(None),
    categoria: Optional[str] = Form(None),
    subcategoria: Optional[str] = Form(None),
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """
    Sube documentos asociados a un cliente existente usando el servicio de versionamiento.
    
    Args:
        files: Lista de archivos a subir
        cliente_id: ID del cliente en PostgreSQL
        tipo_documento: Tipo del documento (opcional)
        categoria: Categoría del documento (opcional)
        subcategoria: Subcategoría del documento (opcional)
    
    Returns:
        Lista de documentos subidos con sus detalles
    """
    if not files:
        raise HTTPException(status_code=400, detail="No se proporcionaron archivos")
    
    current_user = await get_current_user(credentials)
    usuario = None
    if current_user:
        usuario = current_user.get("user_id") or current_user.get("sub") or current_user.get("email")
    
    resultados = []
    errores = []
    
    for file in files:
        try:
            contenido = await file.read()
            nombre_archivo = file.filename or f"documento_{uuid.uuid4().hex[:8]}"
            
            resultado = await documento_versionado_service.subir_documento(
                cliente_id=cliente_id,
                nombre_archivo=nombre_archivo,
                contenido=contenido,
                tipo_documento=tipo_documento,
                categoria=categoria,
                subcategoria=subcategoria,
                usuario=usuario,
                metadata_adicional={
                    "content_type": file.content_type,
                    "origen": "chatbot_archivo",
                    "subido_via": "onboarding"
                }
            )
            
            resultados.append({
                "filename": file.filename,
                "status": resultado.get("status"),
                "mensaje": resultado.get("mensaje"),
                "documento_id": resultado.get("documento_id"),
                "documento_uuid": resultado.get("documento_uuid"),
                "version": resultado.get("version")
            })
            
            logger.info(f"Documento '{file.filename}' subido para cliente {cliente_id}: {resultado.get('status')}")
            
        except Exception as e:
            logger.error(f"Error subiendo documento {file.filename}: {e}")
            errores.append({
                "filename": file.filename,
                "error": str(e)
            })
    
    return {
        "success": len(errores) == 0,
        "cliente_id": cliente_id,
        "documentos_subidos": resultados,
        "errores": errores,
        "total_subidos": len(resultados),
        "total_errores": len(errores)
    }


@router.get("/documentos/{cliente_id}")
async def obtener_documentos_cliente(
    cliente_id: int,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """
    Lista todos los documentos de un cliente.
    """
    try:
        documentos = await documento_versionado_service.listar_documentos_cliente(cliente_id)
        return {
            "success": True,
            "cliente_id": cliente_id,
            "documentos": documentos,
            "total": len(documentos)
        }
    except Exception as e:
        logger.error(f"Error listando documentos para cliente {cliente_id}: {e}")
        raise HTTPException(status_code=503, detail="Unable to retrieve documents. Service temporarily unavailable.")
